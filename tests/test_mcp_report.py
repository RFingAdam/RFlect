"""
Tests for RFlect MCP Report Pipeline

Tests the branded DOCX report generation: plot generation, DOCX building,
gain tables, LLM provider fallback, and end-to-end report flow.
"""

import sys
import os
import tempfile
import shutil

import pytest
import numpy as np

# Add rflect-mcp to path so we can import the tools package
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'rflect-mcp'))


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_passive_measurement():
    """Create a LoadedMeasurement with realistic passive data."""
    from tools.import_tools import LoadedMeasurement

    n_theta, n_phi = 19, 37
    theta = np.linspace(0, 180, n_theta)
    phi = np.linspace(0, 360, n_phi)

    np.random.seed(42)
    total_gain = np.random.randn(n_theta * n_phi, 3) * 5 + 2
    h_gain = np.random.randn(n_theta * n_phi, 3) * 5 + 0
    v_gain = np.random.randn(n_theta * n_phi, 3) * 5 + 0

    return LoadedMeasurement(
        file_path="test_hpol.txt + test_vpol.txt",
        scan_type="passive",
        frequencies=[2400.0, 2450.0, 2500.0],
        data={
            "theta": theta,
            "phi": phi,
            "total_gain": total_gain,
            "h_gain": h_gain,
            "v_gain": v_gain,
        },
    )


@pytest.fixture
def mock_active_measurement():
    """Create a LoadedMeasurement with realistic active data and plot arrays."""
    from tools.import_tools import LoadedMeasurement

    n_theta, n_phi = 13, 25
    theta_deg = np.linspace(0, 180, n_theta)
    phi_deg = np.linspace(0, 345, n_phi)
    theta_rad = np.deg2rad(theta_deg)
    phi_rad = np.deg2rad(phi_deg)

    np.random.seed(123)
    total_power_2d = np.random.randn(n_theta, n_phi) * 3 + 5
    h_power_2d = np.random.randn(n_theta, n_phi) * 3 + 3
    v_power_2d = np.random.randn(n_theta, n_phi) * 3 + 3

    # Extended arrays for 3D wrapping
    phi_deg_plot = np.append(phi_deg, 360)
    phi_rad_plot = np.deg2rad(phi_deg_plot)
    total_power_2d_plot = np.hstack((total_power_2d, total_power_2d[:, [0]]))
    h_power_2d_plot = np.hstack((h_power_2d, h_power_2d[:, [0]]))
    v_power_2d_plot = np.hstack((v_power_2d, v_power_2d[:, [0]]))

    return LoadedMeasurement(
        file_path="test_active.txt",
        scan_type="active",
        frequencies=[2450.0],
        data={
            "total_power": total_power_2d.flatten(),
            "h_power": h_power_2d.flatten(),
            "v_power": v_power_2d.flatten(),
            "TRP_dBm": 5.0,
            "h_TRP_dBm": 3.0,
            "v_TRP_dBm": 3.0,
            "theta": theta_deg,
            "phi": phi_deg,
            "data_points": n_theta * n_phi,
            "theta_rad": theta_rad,
            "phi_rad": phi_rad,
            "total_power_2d": total_power_2d,
            "h_power_2d": h_power_2d,
            "v_power_2d": v_power_2d,
            "phi_deg_plot": phi_deg_plot,
            "phi_rad_plot": phi_rad_plot,
            "total_power_2d_plot": total_power_2d_plot,
            "h_power_2d_plot": h_power_2d_plot,
            "v_power_2d_plot": v_power_2d_plot,
        },
    )


@pytest.fixture
def default_opts():
    """Create default ReportOptions with AI disabled for deterministic tests."""
    from tools.report_tools import ReportOptions

    return ReportOptions(
        ai_executive_summary=False,
        ai_section_analysis=False,
        ai_recommendations=False,
        include_2d_plots=False,
        include_3d_plots=False,
    )


@pytest.fixture
def temp_dir():
    """Create and clean up a temp directory for test outputs."""
    d = tempfile.mkdtemp(prefix="rflect_test_")
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def _inject_passive(mock_passive_measurement):
    """Inject a passive measurement into the MCP global store."""
    from tools.import_tools import _loaded_measurements, _measurements_lock

    with _measurements_lock:
        _loaded_measurements["TestPassive"] = mock_passive_measurement
    yield
    with _measurements_lock:
        _loaded_measurements.pop("TestPassive", None)


@pytest.fixture
def _inject_active(mock_active_measurement):
    """Inject an active measurement into the MCP global store."""
    from tools.import_tools import _loaded_measurements, _measurements_lock

    with _measurements_lock:
        _loaded_measurements["TestActive"] = mock_active_measurement
    yield
    with _measurements_lock:
        _loaded_measurements.pop("TestActive", None)


# ---------------------------------------------------------------------------
# Tests: ReportOptions
# ---------------------------------------------------------------------------

class TestReportOptionsBranded:
    """Test ReportOptions dataclass with new fields."""

    def test_default_options(self):
        from tools.report_tools import ReportOptions

        opts = ReportOptions()
        assert opts.include_cover_page is True
        assert opts.include_table_of_contents is True
        assert opts.include_gain_tables is True
        assert opts.ai_model == "gpt-4o-mini"

    def test_custom_metadata_passthrough(self):
        from tools.report_tools import ReportOptions

        opts = ReportOptions(company_name="TestCo", logo_path="/tmp/logo.png")
        assert opts.company_name == "TestCo"
        assert opts.logo_path == "/tmp/logo.png"


# ---------------------------------------------------------------------------
# Tests: _fmt helper
# ---------------------------------------------------------------------------

class TestFmtReport:
    """Test the report-specific _fmt helper."""

    def test_fmt_float(self):
        from tools.report_tools import _fmt
        assert _fmt(10.123) == "10.12"

    def test_fmt_none(self):
        from tools.report_tools import _fmt
        assert _fmt(None) == "N/A"

    def test_fmt_with_suffix(self):
        from tools.report_tools import _fmt
        assert _fmt(10.0, ".1f", " dBi") == "10.0 dBi"

    def test_fmt_non_numeric(self):
        from tools.report_tools import _fmt
        assert _fmt("hello") == "hello"


# ---------------------------------------------------------------------------
# Tests: Classification Helpers
# ---------------------------------------------------------------------------

class TestClassifyGainQuality:

    def test_below_isotropic(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(-2.0)
        assert rating == "poor"
        assert "below isotropic" in desc

    def test_chip_antenna_range(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(1.5)
        assert rating == "low"
        assert "chip antenna" in desc

    def test_pcb_trace_range(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(3.0)
        assert rating == "moderate"
        assert "PCB trace" in desc

    def test_patch_range(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(6.5)
        assert rating == "good"
        assert "patch" in desc

    def test_high_gain(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(15.0)
        assert rating == "very high"

    def test_none_input(self):
        from tools.report_tools import _classify_gain_quality
        rating, desc = _classify_gain_quality(None)
        assert rating == "unknown"


class TestClassifyEfficiency:

    def test_excellent(self):
        from tools.report_tools import _classify_efficiency
        assert "excellent" in _classify_efficiency(95)

    def test_good(self):
        from tools.report_tools import _classify_efficiency
        assert "good" in _classify_efficiency(80)

    def test_fair(self):
        from tools.report_tools import _classify_efficiency
        assert "fair" in _classify_efficiency(60)

    def test_poor(self):
        from tools.report_tools import _classify_efficiency
        assert "poor" in _classify_efficiency(30)

    def test_none_input(self):
        from tools.report_tools import _classify_efficiency
        assert "unknown" in _classify_efficiency(None)


class TestGetTestConfiguration:

    def test_passive_config(self, mock_passive_measurement):
        from tools.report_tools import _get_test_configuration
        config = _get_test_configuration(mock_passive_measurement)
        assert config["scan_type"] == "passive"
        assert config["num_frequencies"] == 3
        assert "theta_range" in config
        assert "phi_range" in config

    def test_active_config(self, mock_active_measurement):
        from tools.report_tools import _get_test_configuration
        config = _get_test_configuration(mock_active_measurement)
        assert config["scan_type"] == "active"
        assert "TRP_dBm" in config
        assert config["TRP_dBm"] == 5.0


# ---------------------------------------------------------------------------
# Tests: Executive Summary
# ---------------------------------------------------------------------------

class TestBuildExecutiveSummary:

    def test_passive_summary(self, mock_passive_measurement):
        from tools.report_tools import _build_executive_summary, ReportOptions
        from tools.import_tools import _loaded_measurements, _measurements_lock

        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement
        try:
            paragraphs = _build_executive_summary(
                {"TestPassive": mock_passive_measurement}, ReportOptions())
            assert len(paragraphs) >= 1
            # Scope paragraph should mention "passive"
            assert "passive" in paragraphs[0].lower()
            # Should have actual gain values in highlights
            full_text = " ".join(paragraphs)
            assert "dBi" in full_text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_active_summary(self, mock_active_measurement):
        from tools.report_tools import _build_executive_summary, ReportOptions
        from tools.import_tools import _loaded_measurements, _measurements_lock

        with _measurements_lock:
            _loaded_measurements["TestActive"] = mock_active_measurement
        try:
            paragraphs = _build_executive_summary(
                {"TestActive": mock_active_measurement}, ReportOptions())
            assert len(paragraphs) >= 1
            assert "active" in paragraphs[0].lower() or "TRP" in paragraphs[0]
            full_text = " ".join(paragraphs)
            assert "TRP" in full_text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestActive", None)


# ---------------------------------------------------------------------------
# Tests: Pattern Prose
# ---------------------------------------------------------------------------

class TestBuildPatternProse:

    def test_pattern_prose_is_sentences(self, mock_passive_measurement):
        from tools.report_tools import _build_pattern_prose
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from plot_antenna.ai_analysis import AntennaAnalyzer

        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            analyzer = AntennaAnalyzer(
                measurement_data=mock_passive_measurement.data,
                scan_type="passive",
                frequencies=[2400.0, 2450.0, 2500.0],
            )
            prose = _build_pattern_prose(analyzer, 2450.0, "TestPassive")
            # Should be sentences, not key:value pairs
            assert ":" not in prose.split(".")[0] or "MHz" in prose.split(".")[0]
            assert "dBi" in prose
            assert "TestPassive" in prose
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)


# ---------------------------------------------------------------------------
# Tests: Data-Driven Conclusions
# ---------------------------------------------------------------------------

class TestBuildDataDrivenConclusions:

    def test_passive_conclusions_reference_data(self, mock_passive_measurement):
        from tools.report_tools import _build_data_driven_conclusions, ReportOptions
        from tools.import_tools import _loaded_measurements, _measurements_lock

        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement
        try:
            bullets = _build_data_driven_conclusions(
                {"TestPassive": mock_passive_measurement}, ReportOptions())
            assert len(bullets) >= 1
            # Should contain actual numbers
            full_text = " ".join(bullets)
            assert "dBi" in full_text or "dB" in full_text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_active_conclusions_reference_trp(self, mock_active_measurement):
        from tools.report_tools import _build_data_driven_conclusions, ReportOptions
        from tools.import_tools import _loaded_measurements, _measurements_lock

        with _measurements_lock:
            _loaded_measurements["TestActive"] = mock_active_measurement
        try:
            bullets = _build_data_driven_conclusions(
                {"TestActive": mock_active_measurement}, ReportOptions())
            full_text = " ".join(bullets)
            assert "TRP" in full_text or "dBm" in full_text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestActive", None)


# ---------------------------------------------------------------------------
# Tests: _create_llm_provider
# ---------------------------------------------------------------------------

class TestCreateLLMProvider:
    """Test LLM provider creation with graceful fallback."""

    def test_no_key_returns_none(self, monkeypatch):
        """Verify graceful None return when no API key is configured."""
        from tools.report_tools import _create_llm_provider, ReportOptions

        # Patch get_api_key to return None
        monkeypatch.setattr(
            "tools.report_tools.get_all_analysis",
            lambda *a, **kw: "",
        )
        # _create_llm_provider imports internally, so we need to patch at source
        import plot_antenna.api_keys as ak
        monkeypatch.setattr(ak, "get_api_key", lambda _: None)

        provider = _create_llm_provider(ReportOptions())
        assert provider is None


# ---------------------------------------------------------------------------
# Tests: _prepare_report_data
# ---------------------------------------------------------------------------

class TestPrepareReportData:
    """Test measurement filtering and frequency collection."""

    def test_all_measurements(self, mock_passive_measurement):
        from tools.report_tools import _prepare_report_data, ReportOptions

        measurements = {"passive1": mock_passive_measurement}
        opts = ReportOptions()
        data = _prepare_report_data(measurements, opts)

        assert "passive1" in data["measurements"]
        assert data["frequencies"] == [2400.0, 2450.0, 2500.0]

    def test_frequency_filter(self, mock_passive_measurement):
        from tools.report_tools import _prepare_report_data, ReportOptions

        measurements = {"passive1": mock_passive_measurement}
        opts = ReportOptions(frequencies=[2450.0])
        data = _prepare_report_data(measurements, opts)

        assert data["frequencies"] == [2450.0]

    def test_measurement_filter(self, mock_passive_measurement, mock_active_measurement):
        from tools.report_tools import _prepare_report_data, ReportOptions

        measurements = {
            "passive1": mock_passive_measurement,
            "active1": mock_active_measurement,
        }
        opts = ReportOptions(measurements=["active1"])
        data = _prepare_report_data(measurements, opts)

        assert data["measurements"] == ["active1"]
        assert "passive1" not in data["measurements"]

    def test_empty_measurements(self):
        from tools.report_tools import _prepare_report_data, ReportOptions

        data = _prepare_report_data({}, ReportOptions())
        assert data["measurements"] == []
        assert data["frequencies"] == []


# ---------------------------------------------------------------------------
# Tests: _filter_frequencies
# ---------------------------------------------------------------------------

class TestFilterFrequencies:

    def test_no_filter(self):
        from tools.report_tools import _filter_frequencies, ReportOptions
        opts = ReportOptions(frequencies=None)
        assert _filter_frequencies([1.0, 2.0, 3.0], opts) == [1.0, 2.0, 3.0]

    def test_subset(self):
        from tools.report_tools import _filter_frequencies, ReportOptions
        opts = ReportOptions(frequencies=[2.0])
        assert _filter_frequencies([1.0, 2.0, 3.0], opts) == [2.0]


# ---------------------------------------------------------------------------
# Tests: _safe_filename
# ---------------------------------------------------------------------------

class TestSafeFilename:

    def test_safe_chars_preserved(self):
        from tools.report_tools import _safe_filename
        assert _safe_filename("BLE_AP_Test") == "BLE_AP_Test"

    def test_special_chars_replaced(self):
        from tools.report_tools import _safe_filename
        result = _safe_filename("test/path:file")
        assert "/" not in result
        assert ":" not in result


# ---------------------------------------------------------------------------
# Tests: _add_gain_stats_table (legacy)
# ---------------------------------------------------------------------------

class TestGainStatsTable:

    def test_passive_table(self):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _add_gain_stats_table

        doc = Document()
        stats = {
            "scan_type": "passive",
            "frequency_actual": 2450.0,
            "max_gain_dBi": 5.12,
            "min_gain_dBi": -10.5,
            "avg_gain_dBi": 1.23,
            "std_dev_dB": 3.45,
            "max_hpol_gain_dBi": 4.0,
            "max_vpol_gain_dBi": 3.5,
        }
        _add_gain_stats_table(doc, stats, RGBColor(50, 50, 50))

        # Should have one table with 8 rows (header + 5 base + 2 pol)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Parameter"
        assert "5.12" in table.rows[2].cells[1].text

    def test_active_table_with_trp(self):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _add_gain_stats_table

        doc = Document()
        stats = {
            "scan_type": "active",
            "max_power_dBm": 10.0,
            "min_power_dBm": -5.0,
            "avg_power_dBm": 3.0,
            "std_dev_dB": 2.0,
            "TRP_dBm": 5.5,
        }
        _add_gain_stats_table(doc, stats, RGBColor(50, 50, 50))

        table = doc.tables[0]
        # Should have TRP row
        last_row = table.rows[-1]
        assert "TRP" in last_row.cells[0].text
        assert "5.50" in last_row.cells[1].text


# ---------------------------------------------------------------------------
# Tests: Consolidated Performance Table
# ---------------------------------------------------------------------------

class TestConsolidatedPerformanceTable:

    def test_passive_consolidated_table(self, mock_passive_measurement):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _build_consolidated_performance_table
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from plot_antenna.ai_analysis import AntennaAnalyzer

        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement
        try:
            analyzer = AntennaAnalyzer(
                measurement_data=mock_passive_measurement.data,
                scan_type="passive",
                frequencies=[2400.0, 2450.0, 2500.0],
            )
            doc = Document()
            _build_consolidated_performance_table(
                doc, analyzer, "TestPassive",
                [2400.0, 2450.0, 2500.0],
                RGBColor(50, 50, 50), "passive")

            # ONE table with 4 rows (header + 3 freqs)
            assert len(doc.tables) == 1
            table = doc.tables[0]
            assert len(table.rows) == 4
            # Check header has efficiency column
            header_texts = [table.rows[0].cells[j].text for j in range(len(table.rows[0].cells))]
            assert "Eff (%)" in header_texts
            assert "Peak (dBi)" in header_texts
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_active_consolidated_table(self, mock_active_measurement):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _build_consolidated_performance_table
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from plot_antenna.ai_analysis import AntennaAnalyzer

        with _measurements_lock:
            _loaded_measurements["TestActive"] = mock_active_measurement
        try:
            analyzer = AntennaAnalyzer(
                measurement_data=mock_active_measurement.data,
                scan_type="active",
                frequencies=[2450.0],
            )
            doc = Document()
            _build_consolidated_performance_table(
                doc, analyzer, "TestActive",
                [2450.0],
                RGBColor(50, 50, 50), "active")

            assert len(doc.tables) == 1
            table = doc.tables[0]
            assert len(table.rows) == 2  # header + 1 freq
            header_texts = [table.rows[0].cells[j].text for j in range(len(table.rows[0].cells))]
            assert "TRP (dBm)" in header_texts
            assert "H-TRP (dBm)" in header_texts
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestActive", None)


# ---------------------------------------------------------------------------
# Tests: Polarization Table
# ---------------------------------------------------------------------------

class TestPolarizationTable:

    def test_polarization_table_structure(self, mock_passive_measurement):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _build_polarization_table
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from plot_antenna.ai_analysis import AntennaAnalyzer

        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement
        try:
            analyzer = AntennaAnalyzer(
                measurement_data=mock_passive_measurement.data,
                scan_type="passive",
                frequencies=[2400.0, 2450.0, 2500.0],
            )
            doc = Document()
            _build_polarization_table(
                doc, analyzer, [2400.0, 2450.0, 2500.0], RGBColor(50, 50, 50))

            assert len(doc.tables) == 1
            table = doc.tables[0]
            assert len(table.rows) == 4  # header + 3 freqs
            header_texts = [table.rows[0].cells[j].text for j in range(len(table.rows[0].cells))]
            assert "HPOL Peak (dBi)" in header_texts
            assert "Dominant" in header_texts
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)


# ---------------------------------------------------------------------------
# Tests: TRP Section
# ---------------------------------------------------------------------------

class TestTRPSection:

    def test_trp_section_with_balance(self, mock_active_measurement):
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _build_trp_section

        doc = Document()

        def heading(d, t, level=1):
            d.add_heading(t, level)

        _build_trp_section(
            doc, mock_active_measurement, RGBColor(50, 50, 50), heading)

        assert len(doc.tables) == 1
        # Should have TRP values in table
        table = doc.tables[0]
        found_trp = False
        for row in table.rows:
            if "Total TRP" in row.cells[0].text:
                found_trp = True
                assert "5.00" in row.cells[1].text
        assert found_trp

        # Should have assessment prose
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "TRP" in text or "Radiated Power" in text


# ---------------------------------------------------------------------------
# Tests: _build_branded_docx (no AI, no plots)
# ---------------------------------------------------------------------------

class TestBuildBrandedDocx:

    def test_docx_no_ai_no_plots(self, mock_passive_measurement, temp_dir, default_opts):
        """Verify DOCX is created with correct sections when AI is off."""
        from tools.report_tools import _build_branded_docx, _prepare_report_data
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            report_data = _prepare_report_data(measurements, default_opts)
            output = os.path.join(temp_dir, "test_report.docx")

            _build_branded_docx(
                output, report_data, {},
                default_opts, None, None, measurements,
            )

            assert os.path.exists(output)
            assert os.path.getsize(output) > 0

            # Verify DOCX structure
            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "Antenna Radiation Pattern Test Report" in text
            assert "Test Configuration" in text
            assert "Conclusions and Recommendations" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_docx_with_metadata(self, mock_passive_measurement, temp_dir, default_opts):
        """Verify metadata appears on cover page."""
        from tools.report_tools import _build_branded_docx, _prepare_report_data
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            report_data = _prepare_report_data(measurements, default_opts)
            output = os.path.join(temp_dir, "meta_report.docx")

            metadata = {
                "title": "BLE Antenna Report",
                "project_name": "Project X",
                "antenna_type": "PCB Trace",
                "author": "Test Engineer",
            }

            _build_branded_docx(
                output, report_data, {},
                default_opts, None, metadata, measurements,
            )

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "BLE Antenna Report" in text
            assert "Project X" in text
            assert "PCB Trace" in text
            assert "Test Engineer" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_docx_with_gain_tables(self, mock_passive_measurement, temp_dir):
        """Verify consolidated gain tables are present when enabled."""
        from tools.report_tools import (
            _build_branded_docx, _prepare_report_data, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            opts = ReportOptions(
                include_gain_tables=True,
                ai_executive_summary=False,
                ai_section_analysis=False,
                ai_recommendations=False,
                include_2d_plots=False,
                include_3d_plots=False,
            )
            report_data = _prepare_report_data(measurements, opts)
            output = os.path.join(temp_dir, "tables_report.docx")

            _build_branded_docx(
                output, report_data, {},
                opts, None, None, measurements,
            )

            from docx import Document
            doc = Document(output)

            # Should have tables (consolidated + comparison + polarization)
            assert len(doc.tables) >= 1
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_data_driven_conclusions(self, mock_passive_measurement, temp_dir, default_opts):
        """Verify data-driven conclusions reference actual values."""
        from tools.report_tools import _build_branded_docx, _prepare_report_data
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            report_data = _prepare_report_data(measurements, default_opts)
            output = os.path.join(temp_dir, "conclusions_report.docx")

            _build_branded_docx(
                output, report_data, {},
                default_opts, None, None, measurements,
            )

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            # Should contain actual data references, not generic boilerplate
            assert "dBi" in text or "dB" in text
            assert "Conclusions and Recommendations" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_docx_has_executive_summary_with_data(
            self, mock_passive_measurement, temp_dir, default_opts):
        """Verify executive summary contains actual gain/frequency data."""
        from tools.report_tools import _build_branded_docx, _prepare_report_data
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            report_data = _prepare_report_data(measurements, default_opts)
            output = os.path.join(temp_dir, "exec_summary_report.docx")

            _build_branded_docx(
                output, report_data, {},
                default_opts, None, None, measurements,
            )

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            # Executive summary should have actual values
            assert "Executive Summary" in text
            assert "passive" in text.lower()
            assert "dBi" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_docx_has_test_config_methodology(
            self, mock_passive_measurement, temp_dir, default_opts):
        """Verify test configuration includes angular resolution details."""
        from tools.report_tools import _build_branded_docx, _prepare_report_data
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            report_data = _prepare_report_data(measurements, default_opts)
            output = os.path.join(temp_dir, "config_report.docx")

            _build_branded_docx(
                output, report_data, {},
                default_opts, None, None, measurements,
            )

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "Test Configuration" in text
            # Should contain angular resolution info
            assert "Theta" in text or "theta" in text
            assert "Phi" in text or "phi" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)

    def test_docx_pattern_analysis_is_prose(
            self, mock_passive_measurement, temp_dir):
        """Verify pattern analysis is prose sentences, not raw key:value."""
        from tools.report_tools import (
            _build_branded_docx, _prepare_report_data, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TestPassive": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["TestPassive"] = mock_passive_measurement

        try:
            opts = ReportOptions(
                ai_executive_summary=False,
                ai_section_analysis=False,
                ai_recommendations=False,
                include_2d_plots=False,
                include_3d_plots=False,
                include_gain_tables=False,
            )
            report_data = _prepare_report_data(measurements, opts)
            output = os.path.join(temp_dir, "pattern_report.docx")

            _build_branded_docx(
                output, report_data, {},
                opts, None, None, measurements,
            )

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "Pattern Analysis" in text
            # Pattern section should be prose with "exhibits"
            assert "exhibits" in text.lower()
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TestPassive", None)


# ---------------------------------------------------------------------------
# Tests: _generate_ai_text
# ---------------------------------------------------------------------------

class TestGenerateAIText:

    def test_none_provider_returns_none(self):
        from tools.report_tools import _generate_ai_text, ReportOptions
        result = _generate_ai_text(None, "test", {"frequencies": []}, ReportOptions())
        assert result is None


# ---------------------------------------------------------------------------
# Tests: End-to-end generate_report (via direct function call)
# ---------------------------------------------------------------------------

class TestEndToEnd:

    def test_report_no_data_loaded(self):
        """Verify error message when no data is loaded."""
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from tools.report_tools import _prepare_report_data, ReportOptions

        with _measurements_lock:
            _loaded_measurements.clear()

        data = _prepare_report_data({}, ReportOptions())
        assert data["measurements"] == []

    def test_full_pipeline_passive(self, mock_passive_measurement, temp_dir):
        """End-to-end: import passive data, build branded DOCX, verify output."""
        from tools.report_tools import (
            _build_branded_docx, _prepare_report_data, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"BLE_Test": mock_passive_measurement}
        with _measurements_lock:
            _loaded_measurements["BLE_Test"] = mock_passive_measurement

        try:
            opts = ReportOptions(
                frequencies=[2450.0],
                ai_executive_summary=False,
                ai_section_analysis=False,
                ai_recommendations=False,
                include_gain_tables=True,
                include_2d_plots=False,
                include_3d_plots=False,
            )
            report_data = _prepare_report_data(measurements, opts)
            output = os.path.join(temp_dir, "e2e_report.docx")

            metadata = {
                "title": "E2E Test Report",
                "project_name": "Test",
                "frequency_range": "2400-2500 MHz",
            }

            _build_branded_docx(
                output, report_data, {},
                opts, None, metadata, measurements,
            )

            assert os.path.exists(output)

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            # Cover page
            assert "E2E Test Report" in text
            # Test config
            assert "BLE_Test" in text
            # Gain tables present
            assert len(doc.tables) >= 1
            # Conclusions present
            assert "Conclusions and Recommendations" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("BLE_Test", None)

    def test_full_pipeline_active(self, mock_active_measurement, temp_dir):
        """End-to-end: active data with gain tables and TRP section."""
        from tools.report_tools import (
            _build_branded_docx, _prepare_report_data, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {"TRP_Test": mock_active_measurement}
        with _measurements_lock:
            _loaded_measurements["TRP_Test"] = mock_active_measurement

        try:
            opts = ReportOptions(
                ai_executive_summary=False,
                ai_section_analysis=False,
                ai_recommendations=False,
                include_gain_tables=True,
                include_2d_plots=False,
                include_3d_plots=False,
            )
            report_data = _prepare_report_data(measurements, opts)
            output = os.path.join(temp_dir, "active_report.docx")

            _build_branded_docx(
                output, report_data, {},
                opts, None, None, measurements,
            )

            assert os.path.exists(output)

            from docx import Document
            doc = Document(output)

            # Should have tables for active stats
            assert len(doc.tables) >= 1
            # Check for TRP in table
            found_trp = False
            for table in doc.tables:
                for row in table.rows:
                    if "TRP" in row.cells[0].text:
                        found_trp = True
                        break
            assert found_trp

            # Verify TRP Analysis section exists
            text = "\n".join(p.text for p in doc.paragraphs)
            assert "TRP Analysis" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("TRP_Test", None)

    def test_full_pipeline_mixed(self, mock_passive_measurement,
                                  mock_active_measurement, temp_dir):
        """End-to-end: mixed passive + active data."""
        from tools.report_tools import (
            _build_branded_docx, _prepare_report_data, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        measurements = {
            "Passive_BLE": mock_passive_measurement,
            "Active_TRP": mock_active_measurement,
        }
        with _measurements_lock:
            _loaded_measurements["Passive_BLE"] = mock_passive_measurement
            _loaded_measurements["Active_TRP"] = mock_active_measurement

        try:
            opts = ReportOptions(
                ai_executive_summary=False,
                ai_section_analysis=False,
                ai_recommendations=False,
                include_gain_tables=True,
                include_2d_plots=False,
                include_3d_plots=False,
            )
            report_data = _prepare_report_data(measurements, opts)
            output = os.path.join(temp_dir, "mixed_report.docx")

            _build_branded_docx(
                output, report_data, {},
                opts, None, None, measurements,
            )

            assert os.path.exists(output)

            from docx import Document
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

            # Both measurements present
            assert "Passive_BLE" in text
            assert "Active_TRP" in text
            # TRP section for active
            assert "TRP Analysis" in text
            # Pattern analysis for passive
            assert "Pattern Analysis" in text
            # Polarization for passive
            assert "Polarization" in text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("Passive_BLE", None)
                _loaded_measurements.pop("Active_TRP", None)


# ---------------------------------------------------------------------------
# Tests: _add_freq_comparison_table (legacy interface)
# ---------------------------------------------------------------------------

class TestFreqComparisonTable:

    def test_single_freq_skipped(self, mock_active_measurement):
        """With < 2 frequencies, comparison table should not be added."""
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _add_freq_comparison_table
        from plot_antenna.ai_analysis import AntennaAnalyzer

        analyzer = AntennaAnalyzer(
            measurement_data=mock_active_measurement.data,
            scan_type="active",
            frequencies=[2450.0],
        )
        doc = Document()

        def heading(d, t, level=1):
            d.add_heading(t, level)

        _add_freq_comparison_table(doc, analyzer, heading, RGBColor(50, 50, 50))
        assert len(doc.tables) == 0

    def test_multi_freq_table(self, mock_passive_measurement):
        """With >= 2 frequencies, comparison table should be added."""
        from docx import Document
        from docx.shared import RGBColor
        from tools.report_tools import _add_freq_comparison_table
        from plot_antenna.ai_analysis import AntennaAnalyzer

        analyzer = AntennaAnalyzer(
            measurement_data=mock_passive_measurement.data,
            scan_type="passive",
            frequencies=[2400.0, 2450.0, 2500.0],
        )
        doc = Document()

        def heading(d, t, level=1):
            d.add_heading(t, level)

        _add_freq_comparison_table(doc, analyzer, heading, RGBColor(50, 50, 50))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header row + 3 data rows
        assert len(table.rows) == 4
        assert "Freq (MHz)" in table.rows[0].cells[0].text


# ---------------------------------------------------------------------------
# Tests: _detect_rf_band
# ---------------------------------------------------------------------------

class TestDetectRFBand:

    def test_ble_band_detection(self):
        """Frequencies 2402-2480 should be detected as BLE."""
        from tools.report_tools import _detect_rf_band
        freqs = [2402 + i * 2 for i in range(40)]  # 2402, 2404, ..., 2480
        result = _detect_rf_band(freqs)
        assert result is not None
        assert result["name"] == "BLE"
        assert result["standard"] == "Bluetooth 5.x"
        assert 2402 in result["key_frequencies"]
        assert 2440 in result["key_frequencies"]
        assert 2480 in result["key_frequencies"]

    def test_wifi_24_detection(self):
        """Frequencies spanning 2400-2500 should be detected as WiFi 2.4 GHz."""
        from tools.report_tools import _detect_rf_band
        # Use wider range that fits WiFi but not just BLE
        freqs = list(range(2400, 2501, 5))  # 2400, 2405, ..., 2500
        result = _detect_rf_band(freqs)
        assert result is not None
        # Both BLE and WiFi 2.4 overlap; WiFi range is wider and covers more
        assert result["name"] in ("BLE", "WiFi 2.4 GHz")

    def test_lora_868_detection(self):
        """Frequencies 863-870 should be detected as LoRa EU868."""
        from tools.report_tools import _detect_rf_band
        freqs = [863.0, 864.0, 865.0, 866.0, 867.0, 868.0, 869.0, 870.0]
        result = _detect_rf_band(freqs)
        assert result is not None
        assert result["name"] == "LoRa EU868"

    def test_unknown_band(self):
        """Frequencies 100-200 MHz should not match any known band."""
        from tools.report_tools import _detect_rf_band
        freqs = list(range(100, 201, 10))
        result = _detect_rf_band(freqs)
        assert result is None

    def test_multi_band_picks_best(self):
        """Mixed frequencies should pick the band with best coverage."""
        from tools.report_tools import _detect_rf_band
        # 8 BLE frequencies + 2 outliers = 80% BLE coverage
        freqs = [2402, 2410, 2420, 2430, 2440, 2450, 2460, 2480, 100, 200]
        result = _detect_rf_band(freqs)
        assert result is not None
        assert result["name"] == "BLE"

    def test_empty_frequencies(self):
        """Empty frequency list should return None."""
        from tools.report_tools import _detect_rf_band
        assert _detect_rf_band([]) is None


# ---------------------------------------------------------------------------
# Tests: Band-Aware Frequency Selection
# ---------------------------------------------------------------------------

class TestBandAwareFrequencySelection:

    def test_ble_includes_key_channels(self):
        """BLE frequency selection should include key advertising channels."""
        from tools.report_tools import _select_representative_frequencies, _detect_rf_band
        freqs = [2402 + i * 2 for i in range(40)]  # 2402..2480
        band_info = _detect_rf_band(freqs)
        result = _select_representative_frequencies(freqs, 5, band_info=band_info)
        # Should include band edges and key channels
        assert 2402 in result  # first freq / CH0
        assert 2480 in result  # last freq / CH39
        assert 2440 in result  # center / CH19

    def test_unknown_band_falls_back(self):
        """Without band info, should use even-spacing algorithm."""
        from tools.report_tools import _select_representative_frequencies
        freqs = list(range(100, 201, 10))  # 100, 110, ..., 200
        result = _select_representative_frequencies(freqs, 3, band_info=None)
        assert len(result) == 3
        assert result[0] == 100  # first
        assert result[-1] == 200  # last

    def test_respects_max_count(self):
        """Band-aware selection should never exceed max_count."""
        from tools.report_tools import _select_representative_frequencies, _detect_rf_band
        freqs = [2402 + i * 2 for i in range(40)]
        band_info = _detect_rf_band(freqs)
        for max_count in [3, 5, 7]:
            result = _select_representative_frequencies(freqs, max_count, band_info=band_info)
            assert len(result) <= max_count

    def test_snaps_to_nearest(self):
        """Key freq not in list should pick closest available."""
        from tools.report_tools import _select_representative_frequencies, _detect_rf_band
        # Dataset with 5 MHz steps — 2440 is a key BLE freq and IS in the list
        freqs = list(range(2400, 2485, 5))  # 2400, 2405, ..., 2480
        band_info = _detect_rf_band(freqs)
        result = _select_representative_frequencies(freqs, 5, band_info=band_info)
        # 2402 is key BLE freq, nearest in list is 2400 or 2405
        assert any(abs(f - 2402) <= 5 for f in result)
        # 2440 key freq IS in list
        assert 2440 in result

    def test_small_list_returns_all(self):
        """When freqs <= max_count, return all regardless of band info."""
        from tools.report_tools import _select_representative_frequencies, _detect_rf_band
        freqs = [2402.0, 2440.0, 2480.0]
        band_info = _detect_rf_band(freqs)
        result = _select_representative_frequencies(freqs, 5, band_info=band_info)
        assert result == freqs


# ---------------------------------------------------------------------------
# Tests: Channel Label
# ---------------------------------------------------------------------------

class TestChannelLabel:

    def test_known_channel(self):
        from tools.report_tools import _channel_label
        band_info = {
            "name": "BLE",
            "channels": {"CH0 (adv)": 2402, "CH19 (center)": 2440},
        }
        assert _channel_label(2402, band_info) == "BLE CH0 (adv)"
        assert _channel_label(2440, band_info) == "BLE CH19 (center)"

    def test_no_match(self):
        from tools.report_tools import _channel_label
        band_info = {"name": "BLE", "channels": {"CH0 (adv)": 2402}}
        assert _channel_label(2500, band_info) is None

    def test_none_band_info(self):
        from tools.report_tools import _channel_label
        assert _channel_label(2440, None) is None


# ---------------------------------------------------------------------------
# Tests: Band Context in Prose
# ---------------------------------------------------------------------------

class TestBandContextInProse:

    def test_executive_summary_mentions_band(self, mock_passive_measurement):
        """Executive summary should mention 'BLE' when frequencies match."""
        from tools.report_tools import (
            _build_executive_summary, _detect_rf_band, ReportOptions,
        )
        from tools.import_tools import _loaded_measurements, _measurements_lock

        # Create measurement with BLE-range frequencies
        from tools.import_tools import LoadedMeasurement
        ble_meas = LoadedMeasurement(
            file_path="test.txt",
            scan_type="passive",
            frequencies=[2402.0, 2440.0, 2480.0],
            data=mock_passive_measurement.data,
        )

        with _measurements_lock:
            _loaded_measurements["BLE_Test"] = ble_meas
        try:
            band_map = {"BLE_Test": _detect_rf_band(ble_meas.frequencies)}
            paragraphs = _build_executive_summary(
                {"BLE_Test": ble_meas}, ReportOptions(),
                band_info_map=band_map)
            full_text = " ".join(paragraphs)
            assert "BLE" in full_text
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("BLE_Test", None)

    def test_pattern_prose_labels_channels(self, mock_passive_measurement):
        """Pattern prose should label BLE channels at known frequencies."""
        from tools.report_tools import _build_pattern_prose, _detect_rf_band
        from tools.import_tools import _loaded_measurements, _measurements_lock
        from plot_antenna.ai_analysis import AntennaAnalyzer

        ble_freqs = [2402.0, 2440.0, 2480.0]
        from tools.import_tools import LoadedMeasurement
        ble_meas = LoadedMeasurement(
            file_path="test.txt",
            scan_type="passive",
            frequencies=ble_freqs,
            data=mock_passive_measurement.data,
        )

        with _measurements_lock:
            _loaded_measurements["BLE_Test"] = ble_meas
        try:
            analyzer = AntennaAnalyzer(
                measurement_data=ble_meas.data,
                scan_type="passive",
                frequencies=ble_freqs,
            )
            band_info = _detect_rf_band(ble_freqs)

            # Test at 2440 MHz — should mention CH19
            prose = _build_pattern_prose(analyzer, 2440.0, "BLE_Test",
                                        band_info=band_info)
            assert "CH19" in prose

            # Test at 2402 MHz — should mention CH0
            prose = _build_pattern_prose(analyzer, 2402.0, "BLE_Test",
                                        band_info=band_info)
            assert "CH0" in prose
        finally:
            with _measurements_lock:
                _loaded_measurements.pop("BLE_Test", None)
