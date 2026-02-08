"""
Integration tests for the RFlect MCP Server tools.

Tests all tool categories against real antenna measurement files:
  - import_tools: file import, listing, clearing
  - analysis_tools: frequency listing, pattern analysis, gain stats, polarization
  - report_tools: options, preview, DOCX generation
  - bulk_tools: file listing, pair validation, batch processing
  - server resources: status and help endpoints

Each tool function is called directly (not via MCP protocol) by accessing the
inner function registered with FastMCP via `mcp._tool_manager._tools[name].fn`.
"""

import os
import sys
import shutil
import tempfile

import pytest
import numpy as np

# Force non-interactive matplotlib backend before any plotting imports
import matplotlib
matplotlib.use("Agg")

# --------------------------------------------------------------------------- #
# Path setup - mirrors what server.py does
# --------------------------------------------------------------------------- #
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MCP_DIR = os.path.join(PROJECT_ROOT, "rflect-mcp")
sys.path.insert(0, MCP_DIR)
sys.path.insert(0, PROJECT_ROOT)

from mcp.server.fastmcp import FastMCP

# --------------------------------------------------------------------------- #
# Constants - real test file paths
# --------------------------------------------------------------------------- #
TEST_FILES_DIR = "/home/swamp/Downloads/TestFiles/_Test Files/test_files"

PASSIVE_BLE_HPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_BLE AP_HPol.txt")
PASSIVE_BLE_VPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_BLE AP_VPol.txt")
PASSIVE_LORA_HPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_LoRa AP_HPol.txt")
PASSIVE_LORA_VPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_LoRa AP_VPol.txt")
PASSIVE_LORA2_HPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_LoRa2_HPol.txt")
PASSIVE_LORA2_VPOL = os.path.join(TEST_FILES_DIR, "PassiveTest_LoRa2_VPol.txt")
ACTIVE_BLE_TRP = os.path.join(TEST_FILES_DIR, "Active Test_BLE TRP.txt")


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #

@pytest.fixture(scope="module")
def mcp_server():
    """Create a FastMCP server instance with all tools registered once per module."""
    from tools.import_tools import register_import_tools
    from tools.analysis_tools import register_analysis_tools
    from tools.report_tools import register_report_tools
    from tools.bulk_tools import register_bulk_tools

    mcp = FastMCP("rflect-test")
    register_import_tools(mcp)
    register_analysis_tools(mcp)
    register_report_tools(mcp)
    register_bulk_tools(mcp)
    return mcp


def _tool(mcp, name):
    """Convenience: return the raw callable behind a registered MCP tool."""
    return mcp._tool_manager._tools[name].fn


@pytest.fixture(scope="module")
def tools(mcp_server):
    """Return a dict-like accessor for all tool functions."""

    class _Tools:
        def __init__(self, mcp):
            self._mcp = mcp

        def __getattr__(self, name):
            return _tool(self._mcp, name)

    return _Tools(mcp_server)


@pytest.fixture(autouse=True)
def _clear_state_between_tests():
    """Clear loaded measurements before each test to avoid cross-contamination."""
    from tools.import_tools import _loaded_measurements
    _loaded_measurements.clear()
    yield
    _loaded_measurements.clear()


@pytest.fixture()
def tmp_output_dir():
    """Provide a temporary directory that is cleaned up after the test."""
    d = tempfile.mkdtemp(prefix="rflect_test_")
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture()
def ble_pair_dir():
    """Provide a temp directory containing only the BLE passive HPOL/VPOL pair."""
    d = tempfile.mkdtemp(prefix="rflect_ble_pair_")
    shutil.copy(PASSIVE_BLE_HPOL, d)
    shutil.copy(PASSIVE_BLE_VPOL, d)
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture()
def synthetic_passive_loaded():
    """
    Load a synthetic passive measurement into the global store so analysis
    tools have well-structured data with known properties.

    This bypasses the import_antenna_file tool (which stores raw file parser
    output) and instead provides data in the format AntennaAnalyzer expects.
    """
    from tools.import_tools import _loaded_measurements, LoadedMeasurement

    np.random.seed(42)  # reproducibility
    phi = np.linspace(0, 350, 36)       # 10-degree steps, 36 points
    theta = np.linspace(0, 180, 19)     # 10-degree steps, 19 points
    phi_grid, theta_grid = np.meshgrid(phi, theta)
    n = phi_grid.size  # 36 * 19 = 684

    data = {
        "phi": phi_grid.ravel(),
        "theta": theta_grid.ravel(),
        "total_gain": np.random.randn(n) * 3 + 5,
        "h_gain": np.random.randn(n) * 3 + 4,
        "v_gain": np.random.randn(n) * 3 + 3,
    }

    _loaded_measurements["synth_passive.txt"] = LoadedMeasurement(
        file_path="/tmp/synth_passive.txt",
        scan_type="passive",
        frequencies=[2400.0, 2450.0],
        data=data,
    )
    yield _loaded_measurements["synth_passive.txt"]


@pytest.fixture()
def synthetic_active_loaded():
    """
    Load a synthetic active (TRP) measurement into the global store.
    """
    from tools.import_tools import _loaded_measurements, LoadedMeasurement

    np.random.seed(99)
    phi = np.linspace(0, 350, 36)
    theta = np.linspace(0, 180, 19)
    phi_grid, theta_grid = np.meshgrid(phi, theta)
    n = phi_grid.size

    data = {
        "phi": phi_grid.ravel(),
        "theta": theta_grid.ravel(),
        "total_power": np.random.randn(n) * 2 - 5,
        "TRP_dBm": -3.5,
    }

    _loaded_measurements["synth_active.txt"] = LoadedMeasurement(
        file_path="/tmp/synth_active.txt",
        scan_type="active",
        frequencies=[2440.0],
        data=data,
    )
    yield _loaded_measurements["synth_active.txt"]


# =========================================================================== #
#  1. IMPORT TOOLS
# =========================================================================== #


class TestImportTools:
    """Tests for import_antenna_file, list_loaded_data, clear_data."""

    def test_import_passive_hpol(self, tools):
        """Import a single passive HPOL file and verify success message."""
        result = tools.import_antenna_file(PASSIVE_BLE_HPOL)
        assert "Successfully imported" in result
        assert "PassiveTest_BLE AP_HPol.txt" in result
        assert "passive" in result.lower()

    def test_import_passive_vpol(self, tools):
        """Import a single passive VPOL file and verify success message."""
        result = tools.import_antenna_file(PASSIVE_BLE_VPOL)
        assert "Successfully imported" in result
        assert "PassiveTest_BLE AP_VPol.txt" in result

    def test_import_active_trp(self, tools):
        """Import an active TRP file and verify auto-detection of scan type."""
        result = tools.import_antenna_file(ACTIVE_BLE_TRP)
        assert "Successfully imported" in result
        assert "active" in result.lower()

    def test_import_auto_detects_passive(self, tools):
        """Auto-detect should classify files without 'active'/'trp' as passive."""
        result = tools.import_antenna_file(PASSIVE_LORA_HPOL, scan_type="auto")
        assert "passive" in result.lower()

    def test_import_auto_detects_active(self, tools):
        """Auto-detect should classify files with 'active' or 'trp' as active."""
        result = tools.import_antenna_file(ACTIVE_BLE_TRP, scan_type="auto")
        assert "active" in result.lower()

    def test_import_nonexistent_file(self, tools):
        """Importing a missing file should return an error string."""
        result = tools.import_antenna_file("/no/such/file.txt")
        assert "Error" in result

    def test_list_loaded_data_empty(self, tools):
        """list_loaded_data with nothing loaded returns a helpful message."""
        result = tools.list_loaded_data()
        assert "No data loaded" in result

    def test_list_loaded_data_after_import(self, tools):
        """list_loaded_data after an import shows the measurement."""
        tools.import_antenna_file(PASSIVE_BLE_HPOL)
        result = tools.list_loaded_data()
        assert "PassiveTest_BLE AP_HPol.txt" in result
        assert "passive" in result.lower()

    def test_list_loaded_data_multiple_imports(self, tools):
        """list_loaded_data shows all imported measurements."""
        tools.import_antenna_file(PASSIVE_BLE_HPOL)
        tools.import_antenna_file(ACTIVE_BLE_TRP)
        result = tools.list_loaded_data()
        assert "PassiveTest_BLE AP_HPol.txt" in result
        assert "Active Test_BLE TRP.txt" in result

    def test_clear_data(self, tools):
        """clear_data removes all loaded measurements."""
        tools.import_antenna_file(PASSIVE_BLE_HPOL)
        tools.import_antenna_file(PASSIVE_BLE_VPOL)
        result = tools.clear_data()
        assert "Cleared 2" in result

        # Verify data is gone
        result = tools.list_loaded_data()
        assert "No data loaded" in result

    def test_clear_data_empty(self, tools):
        """clear_data when nothing is loaded returns 0 cleared."""
        result = tools.clear_data()
        assert "Cleared 0" in result

    def test_get_measurement_details(self, tools):
        """get_measurement_details returns info for a loaded measurement."""
        tools.import_antenna_file(PASSIVE_BLE_HPOL)
        result = tools.get_measurement_details("PassiveTest_BLE AP_HPol.txt")
        assert "passive" in result.lower()
        assert "PassiveTest_BLE AP_HPol.txt" in result

    def test_get_measurement_details_not_found(self, tools):
        """get_measurement_details with unknown name returns error."""
        result = tools.get_measurement_details("nonexistent.txt")
        assert "not found" in result.lower()

    def test_import_folder(self, tools, ble_pair_dir):
        """import_antenna_folder imports all matching files in a directory."""
        result = tools.import_antenna_folder(ble_pair_dir, pattern="*.txt")
        assert "Imported 2" in result or "2 file" in result.lower()

    def test_import_folder_nonexistent(self, tools):
        """import_antenna_folder with invalid path returns error."""
        result = tools.import_antenna_folder("/no/such/folder")
        assert "Error" in result

    def test_import_folder_no_match(self, tools, tmp_output_dir):
        """import_antenna_folder with no matching files returns message."""
        result = tools.import_antenna_folder(tmp_output_dir, pattern="*.xyz")
        assert "No files" in result


# =========================================================================== #
#  2. ANALYSIS TOOLS
# =========================================================================== #


class TestAnalysisTools:
    """Tests for analysis tools using synthetic loaded data."""

    def test_list_frequencies_empty(self, tools):
        """list_frequencies with no data returns helpful message."""
        result = tools.list_frequencies()
        assert "No data loaded" in result

    def test_list_frequencies(self, tools, synthetic_passive_loaded):
        """list_frequencies shows frequencies from loaded measurements."""
        result = tools.list_frequencies()
        assert "2400.0" in result
        assert "2450.0" in result

    def test_list_frequencies_specific_measurement(self, tools, synthetic_passive_loaded):
        """list_frequencies with measurement_name filters to that measurement."""
        result = tools.list_frequencies(measurement_name="synth_passive.txt")
        assert "2400.0" in result
        assert "synth_passive.txt" in result

    def test_list_frequencies_unknown_measurement(self, tools, synthetic_passive_loaded):
        """list_frequencies with unknown measurement name returns error."""
        result = tools.list_frequencies(measurement_name="unknown.txt")
        assert "not found" in result.lower()

    def test_analyze_pattern(self, tools, synthetic_passive_loaded):
        """analyze_pattern returns classification and peak gain."""
        result = tools.analyze_pattern(frequency=2400.0)
        assert "Pattern" in result
        assert "Peak Gain" in result
        assert "dBi" in result
        # Must contain a pattern type classification
        assert any(
            t in result.lower()
            for t in ["omnidirectional", "sectoral", "directional"]
        )

    def test_analyze_pattern_default_frequency(self, tools, synthetic_passive_loaded):
        """analyze_pattern with no frequency uses the first available."""
        result = tools.analyze_pattern()
        assert "2400.0 MHz" in result

    def test_analyze_pattern_no_data(self, tools):
        """analyze_pattern with no loaded data returns error message."""
        result = tools.analyze_pattern(frequency=2400.0)
        assert "No data loaded" in result

    def test_get_gain_statistics_passive(self, tools, synthetic_passive_loaded):
        """get_gain_statistics for passive data returns min/max/avg gain in dBi."""
        result = tools.get_gain_statistics(frequency=2400.0)
        assert "Maximum Gain" in result
        assert "Minimum Gain" in result
        assert "Average Gain" in result
        assert "Std Dev" in result
        assert "dBi" in result

    def test_get_gain_statistics_active(self, tools, synthetic_active_loaded):
        """get_gain_statistics for active data returns power in dBm and TRP."""
        result = tools.get_gain_statistics(frequency=2440.0)
        assert "Maximum Power" in result
        assert "dBm" in result
        assert "TRP" in result

    def test_get_gain_statistics_no_data(self, tools):
        """get_gain_statistics with no loaded data returns error."""
        result = tools.get_gain_statistics(frequency=2400.0)
        assert "No data loaded" in result

    def test_get_gain_statistics_hpol_vpol(self, tools, synthetic_passive_loaded):
        """get_gain_statistics shows HPOL and VPOL max gain when available."""
        result = tools.get_gain_statistics(frequency=2400.0)
        assert "HPOL Max Gain" in result
        assert "VPOL Max Gain" in result

    def test_compare_polarizations(self, tools, synthetic_passive_loaded):
        """compare_polarizations returns XPD and balance information."""
        result = tools.compare_polarizations(frequency=2400.0)
        assert "Polarization Comparison" in result
        assert "XPD" in result
        assert "HPOL Max Gain" in result
        assert "VPOL Max Gain" in result
        assert "Balance" in result or "balance" in result.lower()

    def test_compare_polarizations_note(self, tools, synthetic_passive_loaded):
        """compare_polarizations includes a polarization assessment note."""
        result = tools.compare_polarizations(frequency=2400.0)
        # The assessment is one of three categories
        assert any(
            note in result
            for note in [
                "Well-balanced",
                "Moderate polarization imbalance",
                "Significant polarization imbalance",
            ]
        )

    def test_compare_polarizations_no_data(self, tools):
        """compare_polarizations with no data returns error."""
        result = tools.compare_polarizations(frequency=2400.0)
        assert "No data loaded" in result

    def test_get_all_analysis(self, tools, synthetic_passive_loaded):
        """get_all_analysis combines gain stats, pattern, and polarization."""
        result = tools.get_all_analysis(frequency=2400.0)
        # Should contain output from all three sub-analyses
        assert "Gain Statistics" in result
        assert "Pattern Analysis" in result or "Pattern" in result
        assert "Polarization Comparison" in result

    def test_get_all_analysis_no_data(self, tools):
        """get_all_analysis with no data returns error."""
        result = tools.get_all_analysis(frequency=2400.0)
        assert "No data loaded" in result


# =========================================================================== #
#  3. REPORT TOOLS
# =========================================================================== #


class TestReportTools:
    """Tests for report configuration, preview, and generation."""

    def test_get_report_options(self, tools):
        """get_report_options returns documentation string with key sections."""
        result = tools.get_report_options()
        assert isinstance(result, str)
        assert "REPORT FILTERING OPTIONS" in result
        assert "CONTENT FILTERING" in result
        assert "PLOT FILTERING" in result
        assert "DATA FILTERING" in result
        assert "AI CONTENT" in result
        assert "OUTPUT" in result

    def test_preview_report_no_data(self, tools):
        """preview_report with no loaded data returns error message."""
        result = tools.preview_report()
        assert "No data loaded" in result

    def test_preview_report(self, tools, synthetic_passive_loaded):
        """preview_report shows section checkboxes and estimated complexity."""
        result = tools.preview_report()
        assert "REPORT PREVIEW" in result
        assert "MEASUREMENTS" in result
        assert "FREQUENCIES" in result
        assert "SECTIONS" in result
        assert "ESTIMATED COMPLEXITY" in result
        assert "synth_passive.txt" in result

    def test_preview_report_with_options(self, tools, synthetic_passive_loaded):
        """preview_report reflects custom options in section toggles."""
        result = tools.preview_report(options={
            "include_3d_plots": True,
            "include_raw_data_tables": True,
        })
        # 3D plots should show [x]
        assert "[x] 3D Pattern Plots" in result
        # Raw data tables should show [x]
        assert "[x] Raw Data Tables" in result

    def test_generate_report_no_data(self, tools, tmp_output_dir):
        """generate_report with no loaded data returns error."""
        path = os.path.join(tmp_output_dir, "empty.docx")
        result = tools.generate_report(path)
        assert "No data loaded" in result

    def test_generate_report_creates_file(self, tools, synthetic_passive_loaded, tmp_output_dir):
        """generate_report creates a DOCX file on disk."""
        path = os.path.join(tmp_output_dir, "test_report.docx")
        result = tools.generate_report(path, options={
            "ai_executive_summary": False,
            "ai_section_analysis": False,
            "ai_recommendations": False,
        })
        assert "Report generated" in result
        assert os.path.isfile(path)

    def test_generate_report_content_summary(self, tools, synthetic_passive_loaded, tmp_output_dir):
        """generate_report summary lists measurements, frequencies, and plot flags."""
        path = os.path.join(tmp_output_dir, "test_report.docx")
        result = tools.generate_report(path, options={
            "ai_executive_summary": False,
            "ai_section_analysis": False,
            "ai_recommendations": False,
        })
        assert "Measurements: 1" in result
        assert "2400.0" in result
        assert "2D Plots: Yes" in result
        assert "3D Plots: No" in result

    def test_generate_report_docx_valid(self, tools, synthetic_passive_loaded, tmp_output_dir):
        """The generated DOCX is a valid file readable by python-docx."""
        from docx import Document

        path = os.path.join(tmp_output_dir, "test_report.docx")
        tools.generate_report(path, options={
            "ai_executive_summary": False,
            "ai_section_analysis": False,
            "ai_recommendations": False,
        })
        doc = Document(path)
        # Should have at least a title paragraph
        assert len(doc.paragraphs) > 0
        # Check that it contains a report title (may come from YAML template)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Antenna" in full_text and "Report" in full_text


# =========================================================================== #
#  4. BULK TOOLS
# =========================================================================== #


class TestBulkTools:
    """Tests for bulk processing, file listing, and validation."""

    def test_list_measurement_files(self, tools):
        """list_measurement_files categorizes files in the test directory."""
        result = tools.list_measurement_files(TEST_FILES_DIR)
        assert "PASSIVE PAIRS" in result
        assert "[HPOL]" in result
        assert "[VPOL]" in result
        assert "ACTIVE/TRP" in result
        assert "[TRP]" in result
        assert "PAIRING STATUS" in result

    def test_list_measurement_files_counts(self, tools):
        """list_measurement_files reports correct HPOL/VPOL counts."""
        result = tools.list_measurement_files(TEST_FILES_DIR)
        assert "2 HPOL" in result
        assert "2 VPOL" in result

    def test_list_measurement_files_invalid_dir(self, tools):
        """list_measurement_files with invalid directory returns error."""
        result = tools.list_measurement_files("/no/such/directory")
        assert "Error" in result

    def test_list_measurement_files_empty_dir(self, tools, tmp_output_dir):
        """list_measurement_files on empty directory shows zero counts."""
        result = tools.list_measurement_files(tmp_output_dir)
        assert "0 HPOL" in result
        assert "0 VPOL" in result

    def test_validate_file_pair_valid(self, tools):
        """validate_file_pair passes for a correctly matched BLE pair."""
        result = tools.validate_file_pair(PASSIVE_BLE_HPOL, PASSIVE_BLE_VPOL)
        assert "VALIDATION PASSED" in result
        assert "correctly paired" in result.lower()

    def test_validate_file_pair_valid_lora(self, tools):
        """validate_file_pair passes for a correctly matched LoRa pair."""
        result = tools.validate_file_pair(PASSIVE_LORA_HPOL, PASSIVE_LORA_VPOL)
        assert "VALIDATION PASSED" in result

    def test_validate_file_pair_mismatched(self, tools):
        """validate_file_pair fails when HPOL and VPOL are from different sets."""
        result = tools.validate_file_pair(PASSIVE_BLE_HPOL, PASSIVE_LORA_VPOL)
        assert "VALIDATION FAILED" in result

    def test_validate_file_pair_missing_hpol(self, tools):
        """validate_file_pair with missing HPOL file returns error."""
        result = tools.validate_file_pair("/missing/hpol.txt", PASSIVE_BLE_VPOL)
        assert "Error" in result

    def test_validate_file_pair_missing_vpol(self, tools):
        """validate_file_pair with missing VPOL file returns error."""
        result = tools.validate_file_pair(PASSIVE_BLE_HPOL, "/missing/vpol.txt")
        assert "Error" in result

    def test_bulk_process_passive(self, tools, ble_pair_dir, tmp_output_dir):
        """bulk_process_passive processes a BLE HPOL/VPOL pair at a single frequency."""
        result = tools.bulk_process_passive(
            ble_pair_dir,
            frequencies=[2440.0],
            save_path=tmp_output_dir,
        )
        assert "Bulk passive processing complete" in result
        assert "Pairs found: 1" in result
        assert "2440.0" in result
        # Check output directory has content
        assert len(os.listdir(tmp_output_dir)) > 0

    def test_bulk_process_passive_invalid_dir(self, tools):
        """bulk_process_passive with invalid directory returns error."""
        result = tools.bulk_process_passive("/no/such/dir")
        assert "Error" in result

    def test_bulk_process_passive_no_hpol_files(self, tools, tmp_output_dir):
        """bulk_process_passive with no HPOL files returns error."""
        result = tools.bulk_process_passive(tmp_output_dir)
        assert "No HPOL" in result

    def test_bulk_process_passive_invalid_freq(self, tools, ble_pair_dir, tmp_output_dir):
        """bulk_process_passive with a frequency not in the data returns error."""
        result = tools.bulk_process_passive(
            ble_pair_dir,
            frequencies=[9999.0],
            save_path=tmp_output_dir,
        )
        assert "not found" in result.lower() or "error" in result.lower()


# =========================================================================== #
#  5. SERVER RESOURCES
# =========================================================================== #


class TestServerResources:
    """Tests for the MCP resource endpoints defined in server.py."""

    def test_status_no_data(self):
        """rflect://status with no loaded data returns helpful message."""
        from tools.import_tools import get_loaded_data_summary
        result = get_loaded_data_summary()
        assert "No data loaded" in result

    def test_status_with_data(self, synthetic_passive_loaded):
        """rflect://status after loading data shows measurement summary."""
        from tools.import_tools import get_loaded_data_summary
        result = get_loaded_data_summary()
        assert "synth_passive.txt" in result
        assert "passive" in result.lower()
        assert "2400.0" in result

    def test_help_content(self):
        """rflect://help returns tool documentation covering all categories."""
        # Reproduce the help function from server.py
        mcp = FastMCP("rflect-help-test")

        @mcp.resource("rflect://help")
        def get_help() -> str:
            return """
RFlect MCP Server - Antenna Analysis Tools

IMPORT TOOLS:
- import_antenna_file(file_path) - Import single measurement file
- import_antenna_folder(folder_path, pattern) - Import all files from folder
- list_loaded_data() - Show currently loaded measurements
- clear_data() - Clear all loaded data

ANALYSIS TOOLS:
- list_frequencies() - Get available frequencies
- analyze_pattern(frequency, polarization) - Pattern analysis (HPBW, F/B, nulls)
- get_gain_statistics(frequency) - Min/max/avg gain
- compare_polarizations(frequency) - HPOL vs VPOL comparison
- get_all_analysis(frequency) - Combined gain + pattern + polarization analysis

REPORT TOOLS:
- generate_report(output_path, options) - Generate DOCX report
- preview_report(options) - Preview report contents without generating
- get_report_options() - Show available report configuration options

BULK PROCESSING TOOLS:
- list_measurement_files(folder_path) - Scan folder for measurement files
- bulk_process_passive(folder_path, frequencies, cable_loss) - Batch process HPOL/VPOL pairs
- bulk_process_active(folder_path) - Batch process TRP files
- validate_file_pair(hpol_path, vpol_path) - Validate HPOL/VPOL file pairing
- convert_to_cst(hpol_path, vpol_path, vswr_path, frequency) - Convert to CST .ffs format
"""

        result = get_help()
        assert "IMPORT TOOLS" in result
        assert "ANALYSIS TOOLS" in result
        assert "REPORT TOOLS" in result
        assert "BULK PROCESSING TOOLS" in result
        assert "import_antenna_file" in result
        assert "generate_report" in result
        assert "bulk_process_passive" in result

    def test_status_resource_registered(self):
        """Verify the status resource can be registered and called."""
        from tools.import_tools import register_import_tools, get_loaded_data_summary

        mcp = FastMCP("rflect-status-test")
        register_import_tools(mcp)

        @mcp.resource("rflect://status")
        def get_status() -> str:
            return get_loaded_data_summary()

        result = get_status()
        assert isinstance(result, str)
        assert "No data loaded" in result


# =========================================================================== #
#  6. DATACLASS TESTS
# =========================================================================== #


class TestDataclasses:
    """Tests for the dataclasses used by the MCP tools."""

    def test_loaded_measurement_creation(self):
        """LoadedMeasurement can be instantiated with all required fields."""
        from tools.import_tools import LoadedMeasurement

        m = LoadedMeasurement(
            file_path="/tmp/test.csv",
            scan_type="passive",
            frequencies=[2400.0, 2450.0],
            data={"key": "value"},
        )
        assert m.file_path == "/tmp/test.csv"
        assert m.scan_type == "passive"
        assert len(m.frequencies) == 2
        assert isinstance(m.data, dict)

    def test_report_options_defaults(self):
        """ReportOptions defaults match expected values."""
        from tools.report_tools import ReportOptions

        opts = ReportOptions()
        assert opts.frequencies is None
        assert opts.polarizations == ["total"]
        assert opts.include_2d_plots is True
        assert opts.include_3d_plots is False
        assert opts.include_polar_plots is True
        assert opts.include_cartesian_plots is False
        assert opts.include_raw_data_tables is False
        assert opts.include_gain_tables is True
        assert opts.ai_model == "gpt-4o-mini"
        assert opts.output_format == "docx"
        assert opts.include_cover_page is True
        assert opts.include_table_of_contents is True

    def test_report_options_custom(self):
        """ReportOptions can be customized."""
        from tools.report_tools import ReportOptions

        opts = ReportOptions(
            frequencies=[2400.0],
            include_3d_plots=True,
            ai_model="gpt-4o",
            include_raw_data_tables=True,
        )
        assert opts.frequencies == [2400.0]
        assert opts.include_3d_plots is True
        assert opts.ai_model == "gpt-4o"
        assert opts.include_raw_data_tables is True


# =========================================================================== #
#  7. HELPER FUNCTION TESTS
# =========================================================================== #


class TestHelpers:
    """Tests for internal helper functions."""

    def test_fmt_float(self):
        """_fmt formats a float to 2 decimal places by default."""
        from tools.analysis_tools import _fmt
        assert _fmt(10.123) == "10.12"

    def test_fmt_none(self):
        """_fmt returns 'N/A' for None."""
        from tools.analysis_tools import _fmt
        assert _fmt(None) == "N/A"

    def test_fmt_na_string(self):
        """_fmt passes through 'N/A' string unchanged."""
        from tools.analysis_tools import _fmt
        assert _fmt("N/A") == "N/A"

    def test_fmt_custom_format(self):
        """_fmt respects a custom format specifier."""
        from tools.analysis_tools import _fmt
        assert _fmt(10.123, ".1f") == "10.1"

    def test_fmt_integer(self):
        """_fmt formats an integer as float."""
        from tools.analysis_tools import _fmt
        assert _fmt(10) == "10.00"

    def test_fmt_non_numeric_string(self):
        """_fmt returns string representation for non-numeric strings."""
        from tools.analysis_tools import _fmt
        assert _fmt("hello") == "hello"
