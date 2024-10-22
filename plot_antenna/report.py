from file_utils import read_passive_file, read_active_file
from calculations import angles_match, calculate_passive_variables, calculate_active_variables
from plotting import plot_passive_3d_component, plot_2d_passive_data, plot_active_2d_data

from docx import Document
from docx.shared import Inches
from tkinter import simpledialog, filedialog, Tk
import os

def generate_report(doc_title, images, save_path):
    # Create a new Document
    doc = Document()
    doc.add_heading(doc_title, 0)

    # Add images to the document
    for img_path in images:
        if img_path is not None and os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
            doc.add_page_break()
        else:
            print(f"Warning: Image path {img_path} does not exist or is invalid.")
    
    # Save the document
    doc.save(os.path.join(save_path, f"{doc_title}.docx"))

def save_to_word_report(selected_frequency, freq_list, scan_type, hpol_path, vpol_path, active_path, cable_loss, datasheet_plots):
    # Initialize the GUI
    root = Tk()
    root.withdraw()  # Hide the main window
    
    # Prompt the user for the project name
    project_name = simpledialog.askstring("Input", "Enter Project Name:")
    
    # Prompt the user to select a directory
    directory = filedialog.askdirectory(title="Select Directory to Save Project")
    
    # Check if user provided a project name and directory
    if not project_name or not directory:
        print("Project name or directory not provided. Exiting...")
        return
    
    # Create the directory structure
    project_path = os.path.join(directory, project_name)
    two_d_data_path = os.path.join(project_path, "2D Plots")
    report_path = os.path.join(project_path, "Report")

    # Create directories if they don't exist
    os.makedirs(two_d_data_path, exist_ok=True)
    os.makedirs(report_path, exist_ok=True)
    
    # List to store the paths of images to be added to the report
    image_paths = []

    if scan_type == "active":
        # Placeholder for future implementation
        return
    elif scan_type == "passive":
        parsed_hpol_data, start_phi_h, stop_phi_h, inc_phi_h, start_theta_h, stop_theta_h, inc_theta_h = read_passive_file(hpol_path)
        hpol_data = parsed_hpol_data

        parsed_vpol_data, start_phi_v, stop_phi_v, inc_phi_v, start_theta_v, stop_theta_v, inc_theta_v = read_passive_file(vpol_path)
        vpol_data = parsed_vpol_data
        
        # Check to see if selected files have mismatched frequency or angle data
        angles_match(start_phi_h, stop_phi_h, inc_phi_h, start_theta_h, stop_theta_h, inc_theta_h,
                     start_phi_v, stop_phi_v, inc_phi_v, start_theta_v, stop_theta_v, inc_theta_v)

        # Call Methods to Calculate Required Variables and Set up variables for plotting
        passive_variables = calculate_passive_variables(hpol_data, vpol_data, cable_loss, start_phi_h, stop_phi_h, inc_phi_h, start_theta_h, stop_theta_h, inc_theta_h, freq_list, selected_frequency)
        theta_angles_deg, phi_angles_deg, v_gain_dB, h_gain_dB, Total_Gain_dB = passive_variables
        
        # Call Method to Plot Passive Data
        plot_2d_passive_data(theta_angles_deg, phi_angles_deg, v_gain_dB, h_gain_dB, Total_Gain_dB, freq_list, selected_frequency, datasheet_plots, save_path=two_d_data_path)

        # Collect the paths of the generated images
        for filename in os.listdir(two_d_data_path):
            if filename.endswith(".png"):
                image_paths.append(os.path.join(two_d_data_path, filename))
        
        # 3D plots
        for gain_type in ["total", "hpol", "vpol"]:
            plot_path = plot_passive_3d_component(theta_angles_deg, phi_angles_deg, v_gain_dB, h_gain_dB, Total_Gain_dB, freq_list, float(selected_frequency), gain_type=gain_type, save_path=two_d_data_path)
            image_paths.append(plot_path)
    
    # Generate the report
    generate_report(project_name, image_paths, report_path)
    print(f"Report saved to {report_path}")

    if __name__ == "__main__":
        save_to_word_report(selected_frequency, freq_list, scan_type, hpol_path, vpol_path, active_path, cable_loss, datasheet_plots)
