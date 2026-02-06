from .calculations import angles_match, calculate_passive_variables, calculate_active_variables
from .plotting import (
    plot_passive_3d_component,
    plot_2d_passive_data,
    plot_active_2d_data,
    plot_active_3d_data,
)
from .file_utils import read_active_file, read_passive_file
from . import config

from tkinter import simpledialog, filedialog, Tk
from docx import Document  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore[import-untyped]

import datetime
import os
import base64

# Import centralized API key management
from .api_keys import get_api_key


def _create_report_provider():
    """Create an LLM provider for report generation based on config."""
    try:
        from .llm_provider import create_provider

        provider_name = config.AI_PROVIDER if hasattr(config, "AI_PROVIDER") else "openai"

        if provider_name == "openai":
            api_key = get_api_key()
            if not api_key:
                return None
            model = config.AI_MODEL if hasattr(config, "AI_MODEL") else "gpt-4o-mini"
            return create_provider("openai", api_key=api_key, model=model)
        elif provider_name == "anthropic":
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                return None
            model = (
                config.AI_ANTHROPIC_MODEL
                if hasattr(config, "AI_ANTHROPIC_MODEL")
                else "claude-sonnet-4-20250514"
            )
            return create_provider("anthropic", api_key=api_key, model=model)
        elif provider_name == "ollama":
            model = config.AI_OLLAMA_MODEL if hasattr(config, "AI_OLLAMA_MODEL") else "llama3.1"
            base_url = (
                config.AI_OLLAMA_URL
                if hasattr(config, "AI_OLLAMA_URL")
                else "http://localhost:11434"
            )
            return create_provider("ollama", model=model, base_url=base_url)
    except Exception as e:
        print(f"[WARNING] Could not create AI provider: {e}")
    return None


# Helper function to encode an image as a base64 string for OpenAI API
def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")


# Helper function to detect measurement type from filename
def detect_measurement_type(filename):
    """
    Detect measurement type from image filename.

    Parameters:
    - filename: Image filename (string)

    Returns:
    - str: Measurement type ('polarization', 'active', 'passive', or None)
    """
    filename_lower = filename.lower()

    if (
        "polar" in filename_lower
        or "ar_" in filename_lower
        or "tilt" in filename_lower
        or "sense" in filename_lower
        or "xpd" in filename_lower
    ):
        return "polarization"
    elif "active" in filename_lower or "trp" in filename_lower or "power" in filename_lower:
        return "active"
    elif "passive" in filename_lower or "gain" in filename_lower:
        return "passive"
    else:
        return None


def extract_frequency_from_filename(filename):
    """
    Extract frequency value from filename.

    Parameters:
    - filename: Image filename (string)

    Returns:
    - str: Frequency string (e.g., '2437.0MHz') or None
    """
    import re

    # Match patterns like: 2437.0MHz, 5510.0_MHz, Freq___2437.0MHz
    patterns = [
        r"(\d+(?:\.\d+)?)\s*MHz",
        r"(\d+(?:\.\d+)?)\s*GHz",
        r"Freq[_\s]+(\d+(?:\.\d+)?)",
    ]

    for pattern in patterns:
        match = re.search(pattern, filename, re.IGNORECASE)
        if match:
            freq = match.group(1)
            if "ghz" in filename.lower():
                return f"{float(freq)*1000:.1f}MHz"
            return f"{freq}MHz"
    return None


def extract_plot_subtype(filename):
    """
    Extract the plot subtype from filename for grouping similar plots.

    Returns one of: '3d_hpol', '3d_vpol', '3d_total', 'elevation', 'azimuth', 'other'
    """
    filename_lower = filename.lower()

    if "3d" in filename_lower:
        if "hpol" in filename_lower or "h_pol" in filename_lower:
            return "3d_hpol"
        elif "vpol" in filename_lower or "v_pol" in filename_lower:
            return "3d_vpol"
        elif "total" in filename_lower:
            return "3d_total"
        return "3d_other"
    elif "elevation" in filename_lower:
        return "elevation"
    elif "azimuth" in filename_lower:
        return "azimuth"
    elif "theta" in filename_lower and "90" in filename_lower:
        return "azimuth"  # Total_Power_in_Theta___90deg is azimuth
    else:
        return "other"


def group_images_for_batch_analysis(image_list):
    """
    Group images by frequency and plot type for efficient batch AI analysis.

    Returns a dict structure:
    {
        '2437.0MHz': {
            '3d_plots': [(img1, paired1), (img2, paired2), ...],  # 3D TRP plots
            '2d_cuts': [(img1, None), ...],  # Elevation/Azimuth cuts
        },
        '5510.0MHz': {...}
    }
    """
    grouped = {}

    for img_tuple in image_list:
        img_path = img_tuple[0]
        filename = os.path.basename(img_path)

        freq = extract_frequency_from_filename(filename)
        subtype = extract_plot_subtype(filename)

        # Use 'unknown' for images without detectable frequency
        freq_key = freq if freq else "unknown"

        if freq_key not in grouped:
            grouped[freq_key] = {"3d_plots": [], "2d_cuts": [], "other": []}

        # Categorize by plot type
        if subtype.startswith("3d"):
            grouped[freq_key]["3d_plots"].append(img_tuple)
        elif subtype in ["elevation", "azimuth"]:
            grouped[freq_key]["2d_cuts"].append(img_tuple)
        else:
            grouped[freq_key]["other"].append(img_tuple)

    return grouped


def deduplicate_images(image_tuples):
    """
    Remove duplicate images based on filename (ignoring path differences).
    Some reports have duplicate images that shouldn't be analyzed twice.
    """
    seen_filenames = set()
    unique_images = []

    for img_tuple in image_tuples:
        filename = os.path.basename(img_tuple[0])
        if filename not in seen_filenames:
            seen_filenames.add(filename)
            unique_images.append(img_tuple)

    return unique_images


class RFAnalyzer:
    def __init__(self, use_ai=False, project_context=None):
        """
        Initialize RF Analyzer with optional AI capabilities.

        Parameters:
        - use_ai: Boolean to enable/disable AI analysis
        - project_context: Dictionary containing project information for enhanced AI prompting
          Example: {
              'antenna_type': 'Patch Antenna',
              'frequency_range': '2.4-2.5 GHz',
              'application': 'WiFi 2.4 GHz',
              'requirements': 'Gain > 5 dBi, VSWR < 2:1'
          }
        """
        self.messages = []
        self._provider = _create_report_provider() if use_ai else None
        self.use_ai = use_ai and self._provider is not None
        self.project_context = project_context or {}
        self.analysis_results = []  # Track all image analyses
        self.measurement_stats = {
            "frequencies": set(),
            "max_gain": None,
            "min_gain": None,
            "trp_values": [],
            "issues_found": [],
            "measurement_types": set(),
        }

    def analyze_image(self, image_path, measurement_type=None, is_paired=False):
        """
        Analyze the image using OpenAI if the AI flag is set, or return a placeholder.

        Parameters:
        - image_path: Path to the image file
        - measurement_type: Type of measurement ('passive', 'active', 'polarization', etc.)
        - is_paired: Whether this is part of a 1of2/2of2 pair (combines both views)
        """
        if self.use_ai:
            analysis = self.send_to_openai(
                image_path, self.project_context, measurement_type, is_paired
            )

            # Store analysis result for later aggregation
            self.analysis_results.append(
                {
                    "image_name": os.path.basename(image_path),
                    "measurement_type": measurement_type,
                    "analysis": analysis,
                }
            )

            # Extract metrics from this analysis
            self._extract_metrics(analysis, image_path)

            return analysis
        else:
            return self.generate_placeholder_caption(image_path, measurement_type)

    def analyze_image_batch(self, image_paths, frequency, batch_type, measurement_type=None):
        """
        Analyze a batch of related images with a single AI call for token efficiency.
        Uses the unified provider abstraction.
        """
        if not self.use_ai or not image_paths or not self._provider:
            return self._generate_batch_placeholder(
                image_paths, frequency, batch_type, measurement_type
            )

        # Limit to first 4 images for batch analysis
        images_to_analyze = image_paths[:4]

        # Build batch-specific prompt
        prompt = self._build_batch_prompt(
            frequency, batch_type, measurement_type, len(images_to_analyze)
        )

        try:
            from .llm_provider import LLMMessage

            max_tokens = config.AI_MAX_TOKENS if hasattr(config, "AI_MAX_TOKENS") else 300
            temperature = config.AI_TEMPERATURE if hasattr(config, "AI_TEMPERATURE") else 0.2

            # Encode all images
            image_b64_list = [encode_image(p) for p in images_to_analyze]

            if self._provider.supports_vision():
                msg = LLMMessage(role="user", content=prompt, images=image_b64_list)
            else:
                msg = LLMMessage(
                    role="user",
                    content=prompt
                    + "\n\n[Batch image analysis not available - model does not support vision]",
                )

            response = self._provider.chat(
                [msg],
                max_tokens=max_tokens,
                temperature=temperature,
            )

            analysis = response.content.strip() if response.content else ""

            if analysis:
                self.analysis_results.append(
                    {
                        "image_name": f"Batch: {frequency} {batch_type}",
                        "measurement_type": measurement_type,
                        "analysis": analysis,
                    }
                )
                self._extract_metrics(analysis, images_to_analyze[0])
                return analysis

            return self._generate_batch_placeholder(
                image_paths, frequency, batch_type, measurement_type
            )

        except Exception as e:
            print(f"Batch analysis error: {e}")
            return self._generate_batch_placeholder(
                image_paths, frequency, batch_type, measurement_type
            )

    def _build_batch_prompt(self, frequency, batch_type, measurement_type, image_count):
        """Build a prompt for batch image analysis."""
        base_prompt = f"""You are an expert RF Engineer. Analyze these {image_count} antenna measurement plots at {frequency} and provide a SINGLE consolidated technical summary.

**Important**: Provide ONE unified analysis covering all images, not separate analyses for each."""

        if batch_type == "3d_plots":
            specific_guidance = """
These are 3D radiation pattern plots showing H-pol, V-pol, and Total TRP.

Provide a unified analysis covering:
1. **Total TRP**: Report total radiated power (dBm) and max EIRP direction
2. **Polarization Breakdown**: H-pol vs V-pol TRP values and dominant polarization
3. **Pattern Characteristics**: Spherical coverage quality, main lobe direction, null regions
4. **Overall Assessment**: Brief statement on radiation quality for the intended application"""

        elif batch_type == "2d_cuts":
            specific_guidance = """
These are 2D radiation pattern cuts (azimuth and elevation planes).

Provide a unified analysis covering:
1. **Power Range**: Max/Min/Avg across all cuts with units (dBm)
2. **Pattern Shape**: Omnidirectional vs directional, symmetry assessment
3. **Coverage Quality**: Null depths, ripple, and uniformity across cuts
4. **Key Observations**: Any notable patterns or concerns"""

        else:
            specific_guidance = """
Analyze all images together and provide:
1. **Key Metrics**: Primary values with units
2. **Pattern Summary**: Overall characteristics
3. **Performance Notes**: Notable observations"""

        response_style = (
            config.AI_RESPONSE_STYLE if hasattr(config, "AI_RESPONSE_STYLE") else "concise"
        )
        max_words = config.AI_MAX_WORDS if hasattr(config, "AI_MAX_WORDS") else 100

        output_format = f"""

### Output Requirements:
- Write {max_words}-{max_words + 50} words total (NOT per image)
- Use precise RF terminology with numerical values and units
- Write in formal third person technical style
- NO separate sections for each image - ONE unified summary

### Output Format:
**Performance Summary at {frequency}**: [3-4 sentences covering all plots with key metrics]

**Key Specifications**:
• TRP/Power: [values]
• Pattern: [description]
• Coverage: [assessment]"""

        return base_prompt + specific_guidance + output_format

    # Legacy batch methods removed — batch analysis now uses unified provider in analyze_image_batch()

    def _generate_batch_placeholder(self, image_paths, frequency, batch_type, measurement_type):
        """Generate placeholder text for batch analysis when AI is disabled."""
        count = len(image_paths)

        if batch_type == "3d_plots":
            return f"""**Performance Summary at {frequency}**: These {count} 3D radiation pattern plots show the spherical power distribution including H-pol, V-pol, and Total TRP components.

**Key Specifications**:
• TRP: See individual plot annotations
• Pattern: 3D spherical visualization
• Coverage: Review plots for null regions and main lobe direction"""

        elif batch_type == "2d_cuts":
            return f"""**Performance Summary at {frequency}**: These {count} 2D radiation pattern cuts show azimuth and elevation plane characteristics.

**Key Specifications**:
• Power Range: See plot annotations for max/min/avg values
• Pattern: Review for omnidirectional vs directional behavior
• Coverage: Assess uniformity across angular cuts"""

        else:
            return f"""**Performance Summary at {frequency}**: {count} measurement plots at this frequency point.

**Key Specifications**:
• Metrics: See individual plot annotations
• Pattern: Review for key characteristics"""

    def generate_placeholder_caption(self, image_path, measurement_type=None):
        """Generate an enhanced caption for the image based on the file name and type."""
        filename = os.path.basename(image_path)

        # Detect measurement type from filename if not provided
        if not measurement_type:
            measurement_type = detect_measurement_type(filename)

        # Build caption based on type
        caption = f"**Figure**: {filename}\n\n"

        if measurement_type == "polarization":
            caption += "**Description**: Polarization analysis showing axial ratio, tilt angle, sense, and cross-polarization discrimination across the measurement sphere.\n\n"
            caption += "**Analysis**: Review the axial ratio values for circular polarization quality (AR < 3 dB for good CP), polarization sense for handedness identification (RHCP/LHCP), and XPD for isolation between orthogonal components."
        elif measurement_type == "active":
            caption += "**Description**: Active power measurement pattern showing radiated power distribution.\n\n"
            caption += "**Analysis**: Examine total radiated power (TRP), H-pol and V-pol components, and coverage uniformity across the measurement sphere."
        elif measurement_type == "passive":
            caption += "**Description**: Passive gain measurement pattern showing directional gain characteristics.\n\n"
            caption += "**Analysis**: Review peak gain, pattern shape, null depths, and polarization purity. Compare H-pol and V-pol performance."
        else:
            caption += "**Description**: Antenna measurement results.\n\n"
            caption += "**Analysis**: Analyze key performance metrics visible in the plot."

        return caption

    def _extract_metrics(self, analysis_text, image_path):
        """Extract key metrics from analysis text to populate measurement_stats."""
        import re

        # Extract frequency information
        freq_matches = re.findall(r"(\d+(?:\.\d+)?)\s*(?:MHz|GHz)", analysis_text, re.IGNORECASE)
        for freq in freq_matches:
            # Normalize to GHz format for consistency
            if "mhz" in analysis_text.lower():
                self.measurement_stats["frequencies"].add(f"{float(freq)/1000:.2f} GHz")
            else:
                self.measurement_stats["frequencies"].add(f"{freq} GHz")

        # Extract TRP values (dBm)
        trp_matches = re.findall(
            r"TRP[:\s]+([+-]?\d+(?:\.\d+)?)\s*dBm", analysis_text, re.IGNORECASE
        )
        for trp in trp_matches:
            self.measurement_stats["trp_values"].append(float(trp))

        # Extract gain values (dBi/dBd)
        gain_matches = re.findall(
            r"(?:gain|peak)[:\s]+([+-]?\d+(?:\.\d+)?)\s*dBi?", analysis_text, re.IGNORECASE
        )
        for gain in gain_matches:
            gain_val = float(gain)
            if (
                self.measurement_stats["max_gain"] is None
                or gain_val > self.measurement_stats["max_gain"]
            ):
                self.measurement_stats["max_gain"] = gain_val
            if (
                self.measurement_stats["min_gain"] is None
                or gain_val < self.measurement_stats["min_gain"]
            ):
                self.measurement_stats["min_gain"] = gain_val

        # Track issues/characteristics mentioned
        keywords = [
            "asymmetric",
            "null",
            "ripple",
            "high AR",
            "poor XPD",
            "low efficiency",
            "variation",
        ]
        for keyword in keywords:
            if keyword.lower() in analysis_text.lower():
                self.measurement_stats["issues_found"].append(keyword)

    def generate_executive_summary(self):
        """Generate an intelligent executive summary based on all analyzed measurements."""
        if not self.analysis_results:
            return "This report presents comprehensive antenna measurement results. Key findings and performance metrics are detailed in the following sections."

        # Count measurement types
        type_counts = {}
        for result in self.analysis_results:
            mtype = result.get("measurement_type") or "general"
            type_counts[mtype] = type_counts.get(mtype, 0) + 1

        # Build summary paragraphs
        paragraphs = []

        # Opening paragraph - scope of testing
        type_desc = []
        if "active" in type_counts:
            type_desc.append(
                f"{type_counts['active']} active (TRP) measurement{'s' if type_counts['active'] > 1 else ''}"
            )
        if "passive" in type_counts:
            type_desc.append(
                f"{type_counts['passive']} passive gain pattern{'s' if type_counts['passive'] > 1 else ''}"
            )
        if "polarization" in type_counts:
            type_desc.append(
                f"{type_counts['polarization']} polarization analysis plot{'s' if type_counts['polarization'] > 1 else ''}"
            )

        scope_text = f"This report documents antenna performance characterization comprising {', '.join(type_desc) if type_desc else 'multiple measurements'}"

        # Frequency information
        if self.measurement_stats["frequencies"]:
            freqs = sorted(list(self.measurement_stats["frequencies"]))
            if len(freqs) == 1:
                scope_text += f" at {freqs[0]}"
            elif len(freqs) <= 3:
                scope_text += f" across {', '.join(freqs)}"
            else:
                scope_text += f" spanning {freqs[0]} to {freqs[-1]}"
        scope_text += "."
        paragraphs.append(scope_text)

        # Performance highlights paragraph
        highlights = []

        # Gain metrics
        if self.measurement_stats["max_gain"] is not None:
            highlights.append(f"peak gain of {self.measurement_stats['max_gain']:.1f} dBi")

        # TRP metrics
        if self.measurement_stats["trp_values"]:
            avg_trp = sum(self.measurement_stats["trp_values"]) / len(
                self.measurement_stats["trp_values"]
            )
            if len(self.measurement_stats["trp_values"]) > 1:
                min_trp = min(self.measurement_stats["trp_values"])
                max_trp = max(self.measurement_stats["trp_values"])
                highlights.append(
                    f"TRP ranging from {min_trp:.1f} to {max_trp:.1f} dBm (average: {avg_trp:.1f} dBm)"
                )
            else:
                highlights.append(f"TRP of {avg_trp:.1f} dBm")

        if highlights:
            paragraphs.append(f"Key performance metrics include {' and '.join(highlights)}.")

        # Observations paragraph
        if self.measurement_stats["issues_found"]:
            unique_issues = list(set(self.measurement_stats["issues_found"]))
            if len(unique_issues) <= 2:
                issues_text = " and ".join(unique_issues)
            else:
                issues_text = ", ".join(unique_issues[:-1]) + f", and {unique_issues[-1]}"
            paragraphs.append(
                f"Analysis identified notable pattern characteristics including {issues_text}. Refer to individual measurement sections for detailed assessment."
            )

        # Closing
        paragraphs.append(
            "Detailed findings and quantitative performance data are presented in the following sections."
        )

        return " ".join(paragraphs)

    def generate_conclusions(self):
        """Generate intelligent conclusions based on analyzed measurements."""
        if not self.analysis_results:
            return [
                "Validate all measured parameters against product specification requirements",
                "Confirm antenna performance meets target application requirements across the operational bandwidth",
                "Document measurement conditions and equipment calibration status for traceability",
            ]

        conclusions = []

        # Frequency coverage assessment
        if self.measurement_stats["frequencies"]:
            num_freqs = len(self.measurement_stats["frequencies"])
            if num_freqs > 1:
                conclusions.append(
                    f"Multi-band characterization completed across {num_freqs} frequency points. Individual band performance should be validated against band-specific requirements."
                )

        # TRP/Power analysis
        if self.measurement_stats["trp_values"]:
            if len(self.measurement_stats["trp_values"]) > 1:
                trp_range = max(self.measurement_stats["trp_values"]) - min(
                    self.measurement_stats["trp_values"]
                )
                if trp_range > 6:
                    conclusions.append(
                        f"TRP variation of {trp_range:.1f} dB observed across frequency bands. Verify compliance with regulatory power limits and system link budget margins."
                    )

        # Gain performance
        if self.measurement_stats["max_gain"] is not None:
            conclusions.append(
                f"Peak gain of {self.measurement_stats['max_gain']:.1f} dBi documented. Confirm alignment with product datasheet specifications and application requirements."
            )

        # Pattern quality observations
        asymmetric_count = sum(
            1 for result in self.analysis_results if "asymmetric" in result["analysis"].lower()
        )
        if asymmetric_count > 0:
            conclusions.append(
                f"Pattern asymmetry noted in {asymmetric_count} measurement(s). Evaluate impact on intended coverage requirements and mounting orientation."
            )

        # Null analysis
        null_keywords = ["deep null", "pronounced null", "significant null", "null depth"]
        null_count = sum(
            1
            for result in self.analysis_results
            if any(word in result["analysis"].lower() for word in null_keywords)
        )
        if null_count > 0:
            conclusions.append(
                f"Notable pattern nulls identified in {null_count} measurement(s). Verify null positions are acceptable for the intended installation configuration."
            )

        # Default conclusions if no specific observations
        if not conclusions:
            conclusions = [
                "Measured radiation characteristics are consistent with expected antenna performance for this design type.",
                "All parameters should be validated against product specification limits before release.",
            ]

        # Standard closing recommendations
        conclusions.append(
            "Cross-reference measurement results with electromagnetic simulation data to validate design correlation."
        )
        conclusions.append(
            "Retain measurement data and chamber calibration records for product qualification documentation."
        )

        return conclusions

    def send_to_openai(
        self, image_path, project_context=None, measurement_type=None, is_paired=False
    ):
        """
        Send image to AI provider for analysis.
        Uses the unified provider abstraction to support OpenAI, Anthropic, and Ollama.
        """
        if not self._provider:
            return self.generate_placeholder_caption(image_path, measurement_type)

        base64_image = encode_image(image_path)
        prompt_text = self._build_prompt(project_context, measurement_type, is_paired)

        max_tokens = config.AI_MAX_TOKENS if hasattr(config, "AI_MAX_TOKENS") else 150
        temperature = config.AI_TEMPERATURE if hasattr(config, "AI_TEMPERATURE") else 0.2

        try:
            from .llm_provider import LLMMessage

            if self._provider.supports_vision():
                msg = LLMMessage(role="user", content=prompt_text, images=[base64_image])
            else:
                # Fallback for non-vision models: text-only prompt
                msg = LLMMessage(
                    role="user",
                    content=prompt_text
                    + "\n\n[Image analysis not available - model does not support vision]",
                )

            response = self._provider.chat(
                [msg],
                max_tokens=max_tokens,
                temperature=temperature,
            )

            reply = response.content.strip() if response.content else ""
            if reply:
                self.messages.append({"role": "assistant", "content": reply})
                return reply

            return self.generate_placeholder_caption(image_path, measurement_type)

        except Exception as e:
            print(f"[WARNING] AI analysis error: {str(e)[:100]}")
            return "**AI Analysis Unavailable** - Error during analysis. Report will use placeholder captions."

    def _build_prompt(self, project_context, measurement_type, is_paired=False):
        """Build the analysis prompt based on context and measurement type."""
        # Build enhanced prompt based on context
        base_prompt = """You are an expert RF Engineer specializing in antenna measurements and design validation. 
Analyze the provided antenna measurement plot and provide a comprehensive technical summary."""

        # Add paired image context if applicable
        if is_paired:
            base_prompt += "\n\n**Note**: This image represents a combined view (front and rear perspectives, 1of2 and 2of2). Analyze the complete 3D radiation pattern shown."

        # Add project context if available
        if project_context:
            context_text = "\n\n### Project Context:\n"
            if "antenna_type" in project_context:
                context_text += f"- **Antenna Type**: {project_context['antenna_type']}\n"
            if "frequency_range" in project_context:
                context_text += f"- **Frequency Range**: {project_context['frequency_range']}\n"
            if "application" in project_context:
                context_text += f"- **Application**: {project_context['application']}\n"
            if "requirements" in project_context:
                context_text += f"- **Key Requirements**: {project_context['requirements']}\n"
            base_prompt += context_text

        # Add measurement-specific guidance (simplified for concise responses)
        response_style = (
            config.AI_RESPONSE_STYLE if hasattr(config, "AI_RESPONSE_STYLE") else "concise"
        )

        if response_style == "concise":
            # Concise analysis prompts - professional technical style
            if measurement_type == "polarization":
                analysis_prompt = """
### Analysis Focus:
1. **Axial Ratio**: Report peak and average AR values (dB) with angular coverage
2. **Polarization Sense**: Identify RHCP/LHCP/Linear dominant regions
3. **XPD**: Cross-polarization discrimination values and uniformity
4. **Quality Assessment**: Circular polarization quality (AR < 3 dB coverage)
"""
            elif measurement_type == "passive":
                analysis_prompt = """
### Analysis Focus:
1. **Gain Performance**: Peak gain (dBi), boresight direction, 3 dB beamwidth
2. **Pattern Shape**: Omni/directional classification, symmetry assessment
3. **Null Analysis**: Null positions, depths, and potential impact
4. **Polarization**: H-pol vs V-pol gain comparison, dominant polarization
"""
            elif measurement_type == "active":
                analysis_prompt = """
### Analysis Focus:
1. **TRP**: Total Radiated Power (dBm) with H-pol/V-pol breakdown
2. **Power Distribution**: Max/min radiation directions, coverage uniformity
3. **Pattern Quality**: Hot spots, nulls, or dead zones identification
4. **Efficiency Indicators**: Power balance and radiated power assessment
"""
            else:
                analysis_prompt = """
### Analysis Focus:
1. **Key Metrics**: Primary performance values (dB/dBi/dBm)
2. **Pattern Analysis**: Shape, symmetry, coverage characteristics
3. **Band Identification**: Infer operational frequency band
4. **Performance Notes**: Any notable characteristics or anomalies
"""
        else:
            # Detailed analysis prompts (original verbose style)
            if measurement_type == "polarization":
                analysis_prompt = """
### Analysis Focus for Polarization Measurements:

1. **Axial Ratio Analysis**:
   - Report AR values at boresight and coverage angles
   - Identify regions of circular vs linear polarization
   - Assess AR < 3 dB coverage area for circular polarization quality

2. **Polarization Sense**:
   - Identify dominant polarization (RHCP/LHCP/Linear)
   - Note any polarization transitions across the pattern

3. **Cross-Polarization Discrimination (XPD)**:
   - Report XPD values and their variation
   - Assess isolation between orthogonal polarizations

4. **Tilt Angle**:
   - Describe polarization ellipse orientation patterns
   - Note any systematic rotation behavior

5. **Performance Assessment**:
   - Evaluate if the antenna meets circular polarization requirements
   - Identify areas for potential improvement
"""
            elif measurement_type == "passive":
                analysis_prompt = """
### Analysis Focus for Passive Gain Measurements:

1. **Gain Metrics**:
   - Report peak gain, minimum gain, and average gain values
   - Identify the direction of maximum radiation (boresight)
   - Note gain at key angles (0°, ±30°, ±45°, ±60°)

2. **Pattern Characteristics**:
   - Describe the radiation pattern shape (omnidirectional, directional, etc.)
   - Identify null positions and their depths
   - Assess pattern symmetry

3. **Polarization Performance**:
   - Compare H-pol vs V-pol gain
   - Calculate cross-polarization ratio
   - Identify dominant polarization

4. **3D Pattern Analysis** (if applicable):
   - Describe spherical coverage
   - Identify back lobe levels
   - Assess pattern distortion or irregularities

5. **Band-Specific Performance**:
   - Infer operational band from frequency
   - Compare performance against typical requirements for that band
"""
            elif measurement_type == "active":
                analysis_prompt = """
### Analysis Focus for Active Power Measurements:

1. **Power Metrics**:
   - Report TRP (Total Radiated Power) in dBm
   - Note H-pol and V-pol TRP components
   - Calculate percentage power distribution

2. **Pattern Analysis**:
   - Identify directions of maximum and minimum radiation
   - Assess pattern uniformity
   - Note any significant nulls or hot spots

3. **Efficiency Assessment**:
   - Compare radiated power to expected input power
   - Identify potential losses or inefficiencies

4. **Coverage Evaluation**:
   - Assess coverage uniformity across the sphere
   - Identify blind spots or weak coverage areas
"""
            else:
                # Generic analysis prompt
                analysis_prompt = """
### Analysis Guidelines:

1. **Key Parameters**: Identify and report min, max, and average values for:
   - Gain (dBi or dBm)
   - Efficiency (if applicable)
   - Directivity
   - Pattern characteristics

2. **Frequency Band**: Infer the operational band from the plot or filename.

3. **Pattern Characteristics**: Describe radiation pattern shape, symmetry, nulls, and lobes.

4. **Performance Assessment**: Evaluate if the results meet typical requirements for the inferred application.
"""

        # Configure output style based on config settings
        max_words = config.AI_MAX_WORDS if hasattr(config, "AI_MAX_WORDS") else 80
        include_recommendations = (
            config.AI_INCLUDE_RECOMMENDATIONS
            if hasattr(config, "AI_INCLUDE_RECOMMENDATIONS")
            else False
        )

        if response_style == "concise":
            output_requirements = f"""

### Output Requirements:
- Provide a professional technical caption in {max_words}-{max_words + 40} words
- Use precise RF engineering terminology
- Include specific numerical values with units (dBi, dBm, degrees)
- Write in third person, present tense, formal technical style
- Do NOT use phrases like "I can see", "appears to show", or qualify image quality
- Focus on factual observations and quantitative data

### Output Format (use this exact structure):
**Performance Summary**: [2-3 sentences describing primary characteristics and key metrics]

**Key Specifications**:
• Peak/Max: [value with units] at [direction/angle if visible]
• Min/Coverage: [value or coverage description]
• Pattern Type: [omnidirectional/directional/hemispherical with brief qualifier]
{("• **Note**: [One critical observation if applicable]" if include_recommendations else "")}
"""
        else:
            # Detailed mode (original verbose style)
            output_requirements = f"""

### Output Requirements:
- Use clear, concise technical language
- Organize findings with Markdown formatting (headers, bullet points, bold for emphasis)
- Include specific numerical values where visible
- Keep the response focused and under {max_words*3} words
- Do NOT include phrases like "I cannot see" or "image quality"

### Output Structure:
**Summary**: One-sentence overview
**Key Findings**: Bullet points of main observations
**Performance Assessment**: Brief evaluation against typical requirements
{("**Recommendations**: Suggested improvements" if include_recommendations else "")}
"""

        return base_prompt + analysis_prompt + output_requirements

    def _handle_api_error(self, result, api_name):
        """Handle API errors with helpful messages."""
        error_msg = result.get("error", {}).get("message", "Unknown error")
        error_type = result.get("error", {}).get("type", "unknown")
        error_code = result.get("error", {}).get("code", "unknown")

        # Provide helpful error messages
        if "quota" in error_msg.lower() or "insufficient_quota" in error_type:
            helpful_msg = (
                f"⚠️ OpenAI {api_name} API Quota Exceeded\n"
                "Your OpenAI account has insufficient credits.\n\n"
                "To fix this:\n"
                "1. Visit: https://platform.openai.com/account/billing\n"
                "2. Add payment method and credits ($5-10 recommended)\n"
                "3. Wait a few minutes for credits to activate\n"
                "4. Try generating the report again\n\n"
                "Note: The report will still be generated with placeholder captions."
            )
            print(f"\n{'='*60}")
            print(helpful_msg)
            print(f"{'='*60}\n")
            return "**AI Analysis Unavailable** - OpenAI quota exceeded. Please add billing credits at platform.openai.com/account/billing"
        elif "api_key" in error_msg.lower() or "authentication" in error_msg.lower():
            helpful_msg = (
                f"⚠️ OpenAI {api_name} API Key Issue\n"
                "The API key is invalid or not authorized.\n\n"
                "To fix this:\n"
                "1. Visit: https://platform.openai.com/api-keys\n"
                "2. Generate a new API key\n"
                "3. Update 'openapi.env' file: OPENAI_API_KEY=sk-...\n"
                "4. Restart the application\n"
            )
            print(f"\n{'='*60}")
            print(helpful_msg)
            print(f"{'='*60}\n")
            return "**AI Analysis Unavailable** - Invalid API key. Check openapi.env file."
        elif "model" in error_msg.lower() and (
            "not found" in error_msg.lower() or "does not exist" in error_msg.lower()
        ):
            helpful_msg = (
                f"⚠️ Model Not Available\n"
                f"The model specified in config is not available or not supported by {api_name} API.\n\n"
                f"Error: {error_msg}\n\n"
                "To fix this:\n"
                "1. Check AI_MODEL setting in config_local.py or config_template.py\n"
                "2. Verify model name matches OpenAI's current offerings\n"
                "3. GPT-5 models require Responses API access\n"
                "4. Try switching to 'gpt-4o-mini' for guaranteed compatibility\n"
            )
            print(f"\n{'='*60}")
            print(helpful_msg)
            print(f"{'='*60}\n")
            return f"**AI Analysis Unavailable** - Model not available. Check AI_MODEL in config. Error: {error_msg}"
        else:
            print(f"OpenAI {api_name} API Error ({error_type}): {error_msg}")
            return f"**Analysis unavailable**: {error_msg}"


def generate_report(
    doc_title,
    images,
    save_path,
    analyzer,
    logo_path=None,
    metadata=None,
    include_toc=True,
    include_summary=True,
):
    """
    Generate a comprehensive Word document report with antenna measurement results.
    Uses professional Ezurio branding and formatting.

    Parameters:
    - doc_title: Title of the report
    - images: List of image file paths to include
    - save_path: Directory path where the report will be saved
    - analyzer: RFAnalyzer instance for generating captions
    - logo_path: Optional path to company logo (auto-detects from config if not provided)
    - metadata: Dictionary with report metadata (author, date, project info, etc.)
    - include_toc: Boolean to include table of contents
    - include_summary: Boolean to include executive summary
    """

    # Load branding configuration from config.py
    # Use config values if available, otherwise fall back to defaults
    BRAND_PRIMARY = (
        RGBColor(*config.BRAND_PRIMARY_COLOR)
        if config.BRAND_PRIMARY_COLOR
        else RGBColor(70, 130, 180)
    )
    BRAND_DARK = (
        RGBColor(*config.BRAND_DARK_COLOR) if config.BRAND_DARK_COLOR else RGBColor(50, 50, 50)
    )
    BRAND_LIGHT = (
        RGBColor(*config.BRAND_LIGHT_COLOR) if config.BRAND_LIGHT_COLOR else RGBColor(128, 128, 128)
    )

    brand_name = config.BRAND_NAME if hasattr(config, "BRAND_NAME") else None
    brand_tagline = config.BRAND_TAGLINE if hasattr(config, "BRAND_TAGLINE") else None
    brand_website = config.BRAND_WEBSITE if hasattr(config, "BRAND_WEBSITE") else None
    report_subtitle = (
        config.REPORT_SUBTITLE
        if hasattr(config, "REPORT_SUBTITLE")
        else "Antenna Measurement & Analysis Report"
    )

    # Helper function to add branded headings
    def add_branded_heading(doc, text, level=1):
        """Add a heading with configured branding colors."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                run.font.color.rgb = BRAND_DARK
                run.font.size = Pt(
                    config.HEADING1_FONT_SIZE if hasattr(config, "HEADING1_FONT_SIZE") else 18
                )
            else:
                run.font.color.rgb = BRAND_DARK
                run.font.size = Pt(
                    config.HEADING2_FONT_SIZE if hasattr(config, "HEADING2_FONT_SIZE") else 14
                )
            run.font.bold = True
        return heading

    brand_info = f" with {brand_name} Branding" if brand_name else ""
    print(f"Starting Enhanced Report Generation{brand_info}...")
    doc = Document()

    # Set document margins for professional layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Try to find logo if not provided
    if not logo_path and hasattr(config, "LOGO_FILENAME") and config.LOGO_FILENAME:
        # Look for logo in assets folder
        possible_logo_paths = [
            os.path.join(
                os.path.dirname(os.path.dirname(__file__)), "assets", config.LOGO_FILENAME
            ),
            os.path.join(os.path.dirname(__file__), "..", "assets", config.LOGO_FILENAME),
        ]
        for path in possible_logo_paths:
            if os.path.exists(path):
                logo_path = path
                break

    # Add logo to header
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_run = header_para.add_run()

        logo_width = config.LOGO_WIDTH_INCHES if hasattr(config, "LOGO_WIDTH_INCHES") else 2.0
        header_run.add_picture(logo_path, width=Inches(logo_width))

        # Set logo alignment based on config
        logo_align = config.LOGO_ALIGNMENT if hasattr(config, "LOGO_ALIGNMENT") else "LEFT"
        if logo_align == "CENTER":
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif logo_align == "RIGHT":
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add brand color accent line (simulated with text)
        header_line = header.add_paragraph("_" * 80)
        header_line.runs[0].font.color.rgb = BRAND_PRIMARY
        header_line.runs[0].font.size = Pt(1)
        header_line.paragraph_format.space_before = Pt(3)
        header_line.paragraph_format.space_after = Pt(0)

    # Title page with branded styling
    # Add spacing at top
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = BRAND_DARK
    title_run.font.size = Pt(config.TITLE_FONT_SIZE if hasattr(config, "TITLE_FONT_SIZE") else 28)
    title_run.font.bold = True

    # Subtitle line (if configured)
    if report_subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle_para.add_run(report_subtitle)
        subtitle_run.font.color.rgb = BRAND_LIGHT
        subtitle_run.font.size = Pt(
            config.SUBTITLE_FONT_SIZE if hasattr(config, "SUBTITLE_FONT_SIZE") else 14
        )
        subtitle_run.italic = True

    # Add metadata section with branded styling
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    if metadata:
        meta_table_data = []
        if "project_name" in metadata and metadata["project_name"]:
            meta_table_data.append(("Project:", metadata["project_name"]))
        if "antenna_type" in metadata and metadata["antenna_type"]:
            meta_table_data.append(("Antenna Type:", metadata["antenna_type"]))
        if "frequency_range" in metadata and metadata["frequency_range"]:
            meta_table_data.append(("Frequency Range:", metadata["frequency_range"]))
        if "date" in metadata:
            meta_table_data.append(("Date:", metadata["date"]))
        else:
            meta_table_data.append(("Date:", datetime.datetime.now().strftime("%B %d, %Y")))
        if "author" in metadata and metadata["author"]:
            meta_table_data.append(("Prepared by:", metadata["author"]))

        # Create a clean metadata display
        for label, value in meta_table_data:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            label_run = meta_para.add_run(f"{label} ")
            label_run.font.bold = True
            label_run.font.color.rgb = BRAND_DARK
            label_run.font.size = Pt(12)
            value_run = meta_para.add_run(value)
            value_run.font.color.rgb = BRAND_LIGHT
            value_run.font.size = Pt(12)
            meta_para.paragraph_format.space_after = Pt(3)
    else:
        # Default date with branded styling
        date_para = doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
        date_run.font.color.rgb = BRAND_LIGHT
        date_run.font.size = Pt(12)

    # Add brand footer to title page (if configured)
    if brand_website:
        doc.add_paragraph()
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run(brand_website)
        footer_run.font.color.rgb = BRAND_LIGHT
        footer_run.font.size = Pt(10)
        footer_run.italic = True

    doc.add_page_break()

    # Table of Contents placeholder (if requested)
    if include_toc:
        add_branded_heading(doc, "Table of Contents", level=1)
        toc_para = doc.add_paragraph(
            "< Table of Contents will be generated when document is opened in Word >"
        )
        toc_para.paragraph_format.space_before = Pt(6)
        doc.add_page_break()

    # Executive Summary (if requested and AI is enabled)
    if include_summary and analyzer.use_ai:
        add_branded_heading(doc, "Executive Summary", level=1)

        # Generate intelligent summary based on all analyses
        summary_text = analyzer.generate_executive_summary()

        summary_para = doc.add_paragraph(summary_text)
        summary_para.paragraph_format.space_before = Pt(6)
        doc.add_page_break()

    # Measurement Results Section
    add_branded_heading(doc, "Measurement Results", level=1)

    # Detect 1of2/2of2 paired images - track as tuples (img1, img2 or None)
    image_list = []  # List of tuples: (primary_image, paired_image_or_None)
    paired_images = {}  # Track pairs: base_name -> [1of2_path, 2of2_path]
    processed_images = set()  # Track which images we've already added

    for img_path in images:
        if img_path is None or not os.path.exists(img_path):
            print(f"Warning: Image path {img_path} does not exist or is invalid.")
            continue

        filename = os.path.basename(img_path)

        # Check if this is part of a 1of2/2of2 pair
        if "_1of2" in filename:
            base_name = filename.replace("_1of2", "")
            if base_name not in paired_images:
                paired_images[base_name] = [None, None]
            paired_images[base_name][0] = img_path
        elif "_2of2" in filename:
            base_name = filename.replace("_2of2", "")
            if base_name not in paired_images:
                paired_images[base_name] = [None, None]
            paired_images[base_name][1] = img_path

    # Build image list with pairs grouped together
    for img_path in images:
        if img_path is None or not os.path.exists(img_path):
            continue

        if img_path in processed_images:
            continue

        filename = os.path.basename(img_path)

        # Check if this is part of a pair
        if "_1of2" in filename:
            base_name = filename.replace("_1of2", "")
            img1, img2 = paired_images.get(base_name, (None, None))
            if img1 and img2:
                image_list.append((img1, img2))  # Tuple with both images
                processed_images.add(img1)
                processed_images.add(img2)
            elif img1:
                image_list.append((img1, None))  # Only 1of2 exists
                processed_images.add(img1)
        elif "_2of2" in filename:
            base_name = filename.replace("_2of2", "")
            img1, img2 = paired_images.get(base_name, (None, None))
            if not img1 and img2:  # Only 2of2 exists (1of2 not found)
                image_list.append((img2, None))
                processed_images.add(img2)
            # If both exist, already handled in 1of2 case above
        else:
            # Regular unpaired image
            image_list.append((img_path, None))
            processed_images.add(img_path)

    # Deduplicate images (remove duplicates with same filename)
    image_list = deduplicate_images(image_list)

    # Group images by type for better organization
    image_groups = {"polarization": [], "active": [], "passive": [], "other": []}

    for img_tuple in image_list:
        img_path = img_tuple[0]  # Use primary image for classification

        filename = os.path.basename(img_path)
        measurement_type = detect_measurement_type(filename)

        if measurement_type == "polarization":
            image_groups["polarization"].append(img_tuple)
        elif measurement_type == "active":
            image_groups["active"].append(img_tuple)
        elif measurement_type == "passive":
            image_groups["passive"].append(img_tuple)
        else:
            image_groups["other"].append(img_tuple)

    # Process each group
    figure_num = 1

    for group_name, group_images in image_groups.items():
        if not group_images:
            continue

        # Add section header with branded styling
        if group_name == "polarization":
            add_branded_heading(doc, "Polarization Analysis", level=2)
            intro_para = doc.add_paragraph(
                "The following plots show polarization characteristics including axial ratio, "
                "tilt angle, polarization sense (RHCP/LHCP), and cross-polarization discrimination."
            )
            intro_para.paragraph_format.space_before = Pt(6)
            intro_para.paragraph_format.space_after = Pt(12)
        elif group_name == "active":
            add_branded_heading(doc, "Active Power Measurements", level=2)
            intro_para = doc.add_paragraph(
                "Active measurements showing total radiated power (TRP) and power distribution "
                "across H-pol and V-pol components."
            )
            intro_para.paragraph_format.space_before = Pt(6)
            intro_para.paragraph_format.space_after = Pt(12)
        elif group_name == "passive":
            add_branded_heading(doc, "Passive Gain Measurements", level=2)
            intro_para = doc.add_paragraph(
                "Passive gain patterns showing directional gain characteristics, pattern shape, "
                "and polarization performance."
            )
            intro_para.paragraph_format.space_before = Pt(6)
            intro_para.paragraph_format.space_after = Pt(12)

        # Group images by frequency for batch analysis (token optimization)
        freq_grouped = group_images_for_batch_analysis(group_images)

        # Sort frequencies for consistent ordering
        sorted_freqs = sorted(
            freq_grouped.keys(),
            key=lambda x: (
                float(x.replace("MHz", "").replace("GHz", "000")) if x != "unknown" else 0
            ),
        )

        for freq in sorted_freqs:
            freq_data = freq_grouped[freq]

            # Add frequency subsection header if multiple frequencies
            if len(sorted_freqs) > 1 and freq != "unknown":
                freq_heading = doc.add_paragraph()
                freq_run = freq_heading.add_run(f"Measurements at {freq}")
                freq_run.bold = True
                freq_run.font.color.rgb = BRAND_DARK
                freq_run.font.size = Pt(12)
                freq_heading.paragraph_format.space_before = Pt(12)
                freq_heading.paragraph_format.space_after = Pt(6)

            # Process 3D plots as a batch if there are multiple
            if freq_data["3d_plots"]:
                three_d_images = freq_data["3d_plots"]

                # Add all 3D images first
                for img_tuple in three_d_images:
                    img_path = img_tuple[0]
                    paired_img = img_tuple[1]
                    filename = os.path.basename(img_path)
                    is_paired = paired_img is not None

                    doc.add_picture(img_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    caption_para = doc.add_paragraph()
                    caption_run = caption_para.add_run(f"Figure {figure_num}: ")
                    caption_run.bold = True
                    caption_run.font.color.rgb = BRAND_DARK
                    caption_run.font.size = Pt(11)

                    display_name = (
                        filename.replace("_1of2", " (View 1 of 2)").replace(
                            "_2of2", " (View 1 of 2)"
                        )
                        if is_paired
                        else filename
                    )
                    filename_run = caption_para.add_run(display_name)
                    filename_run.font.color.rgb = BRAND_LIGHT
                    filename_run.font.size = Pt(11)
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    figure_num += 1

                    if is_paired and paired_img:
                        doc.add_picture(paired_img, width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        caption_para2 = doc.add_paragraph()
                        caption_run2 = caption_para2.add_run(f"Figure {figure_num}: ")
                        caption_run2.bold = True
                        caption_run2.font.color.rgb = BRAND_DARK
                        caption_run2.font.size = Pt(11)

                        paired_filename = os.path.basename(paired_img)
                        display_name2 = paired_filename.replace("_2of2", " (View 2 of 2)").replace(
                            "_1of2", " (View 2 of 2)"
                        )
                        filename_run2 = caption_para2.add_run(display_name2)
                        filename_run2.font.color.rgb = BRAND_LIGHT
                        filename_run2.font.size = Pt(11)
                        caption_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        figure_num += 1

                # Generate ONE batch analysis for all 3D plots at this frequency
                three_d_paths = [t[0] for t in three_d_images]
                batch_analysis = analyzer.analyze_image_batch(
                    three_d_paths,
                    freq,
                    "3d_plots",
                    measurement_type=group_name if group_name != "other" else None,
                )
                analysis_para = doc.add_paragraph(batch_analysis)
                analysis_para.paragraph_format.space_before = Pt(6)
                analysis_para.paragraph_format.space_after = Pt(12)

            # Process 2D cuts as a batch if there are multiple
            if freq_data["2d_cuts"]:
                two_d_images = freq_data["2d_cuts"]

                # Add all 2D images first
                for img_tuple in two_d_images:
                    img_path = img_tuple[0]
                    filename = os.path.basename(img_path)

                    doc.add_picture(img_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    caption_para = doc.add_paragraph()
                    caption_run = caption_para.add_run(f"Figure {figure_num}: ")
                    caption_run.bold = True
                    caption_run.font.color.rgb = BRAND_DARK
                    caption_run.font.size = Pt(11)

                    filename_run = caption_para.add_run(filename)
                    filename_run.font.color.rgb = BRAND_LIGHT
                    filename_run.font.size = Pt(11)
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    figure_num += 1

                # Generate ONE batch analysis for all 2D cuts at this frequency
                two_d_paths = [t[0] for t in two_d_images]
                batch_analysis = analyzer.analyze_image_batch(
                    two_d_paths,
                    freq,
                    "2d_cuts",
                    measurement_type=group_name if group_name != "other" else None,
                )
                analysis_para = doc.add_paragraph(batch_analysis)
                analysis_para.paragraph_format.space_before = Pt(6)
                analysis_para.paragraph_format.space_after = Pt(12)

            # Process 'other' images individually (fallback for unrecognized types)
            for img_tuple in freq_data["other"]:
                img_path = img_tuple[0]
                paired_img = img_tuple[1]
                filename = os.path.basename(img_path)
                is_paired = paired_img is not None
                measurement_type = group_name if group_name != "other" else None

                # Individual analysis for unrecognized images
                analysis = analyzer.analyze_image(img_path, measurement_type, is_paired)

                doc.add_picture(img_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Figure {figure_num}: ")
                caption_run.bold = True
                caption_run.font.color.rgb = BRAND_DARK
                caption_run.font.size = Pt(11)

                display_name = (
                    filename.replace("_1of2", " (View 1 of 2)").replace("_2of2", " (View 1 of 2)")
                    if is_paired
                    else filename
                )
                filename_run = caption_para.add_run(display_name)
                filename_run.font.color.rgb = BRAND_LIGHT
                filename_run.font.size = Pt(11)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                figure_num += 1

                if is_paired and paired_img:
                    doc.add_picture(paired_img, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    caption_para2 = doc.add_paragraph()
                    caption_run2 = caption_para2.add_run(f"Figure {figure_num}: ")
                    caption_run2.bold = True
                    caption_run2.font.color.rgb = BRAND_DARK
                    caption_run2.font.size = Pt(11)

                    paired_filename = os.path.basename(paired_img)
                    display_name2 = paired_filename.replace("_2of2", " (View 2 of 2)").replace(
                        "_1of2", " (View 2 of 2)"
                    )
                    filename_run2 = caption_para2.add_run(display_name2)
                    filename_run2.font.color.rgb = BRAND_LIGHT
                    filename_run2.font.size = Pt(11)
                    caption_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    figure_num += 1

                analysis_para = doc.add_paragraph(analysis)
                analysis_para.paragraph_format.space_before = Pt(6)
                analysis_para.paragraph_format.space_after = Pt(12)

    # Add conclusions/recommendations section
    doc.add_page_break()
    add_branded_heading(doc, "Conclusions and Recommendations", level=1)

    # Generate intelligent conclusions based on measurements
    if analyzer.use_ai:
        conclusions_list = analyzer.generate_conclusions()
        doc.add_paragraph("Based on the measurement results presented in this report:")
        for conclusion in conclusions_list:
            doc.add_paragraph(f"• {conclusion}", style="List Bullet")
    else:
        doc.add_paragraph("Based on the measurement results presented in this report:")
        doc.add_paragraph(
            "• Review all performance metrics against specification requirements",
            style="List Bullet",
        )
        doc.add_paragraph(
            "• Verify antenna performance meets application needs across the operational bandwidth",
            style="List Bullet",
        )
        doc.add_paragraph(
            "• Consider additional measurements or design iterations if performance gaps are identified",
            style="List Bullet",
        )

    # Add brand contact footer (if configured)
    if brand_tagline or brand_website:
        doc.add_paragraph()
        doc.add_paragraph()

        if brand_tagline and brand_name:
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            footer_run = footer_para.add_run(f"{brand_name} | {brand_tagline}")
            footer_run.font.color.rgb = BRAND_LIGHT
            footer_run.font.size = Pt(10)
            footer_run.italic = True

        if brand_website:
            footer_url = doc.add_paragraph()
            footer_url.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            url_run = footer_url.add_run(brand_website)
            url_run.font.color.rgb = BRAND_PRIMARY
            url_run.font.size = Pt(10)
            url_run.bold = True

    # Save the final report
    output_filename = os.path.join(save_path, f"{doc_title}.docx")
    doc.save(output_filename)
    print(f"✓ Report Generation Complete{brand_info}!")
    print(f"  Saved to: {output_filename}")
    return output_filename


def save_to_results_folder(
    selected_frequency,
    freq_list,
    scan_type,
    hpol_path,
    vpol_path,
    active_path,
    cable_loss,
    datasheet_plots,
    axis_mode,
    zmin,
    zmax,
    word=False,
    logo_path=None,
):  # Initialize the GUI
    root = Tk()
    root.withdraw()  # Hide the main window

    # Prompt the user for the project name
    project_name = simpledialog.askstring("Input", "Enter Project Name:")

    # Prompt the user to select a directory
    directory = filedialog.askdirectory(title="Select Directory to Save Project")

    # Check if user provided a project name and directory
    if not project_name or not directory:
        print("Project name or directory not provided. Exiting...")
        return

    # Create the directory structure
    project_path = os.path.join(directory, project_name)
    two_d_data_path = os.path.join(project_path, "2D Plots")
    if word:
        report_path = os.path.join(project_path, "Report")
        os.makedirs(report_path, exist_ok=True)
    if scan_type == "active":
        three_d_data_path = os.path.join(project_path, "3D Plots")
        os.makedirs(three_d_data_path, exist_ok=True)
    elif scan_type == "passive":
        user_selected_frequency_folder_name = f"3D Plots at {selected_frequency} MHz"
        user_selected_frequency_path = os.path.join(
            project_path, user_selected_frequency_folder_name
        )
        os.makedirs(user_selected_frequency_path, exist_ok=True)
    os.makedirs(two_d_data_path, exist_ok=True)

    # Call the modified plotting functions with save_path argument to save the plots
    if scan_type == "active":
        # Perform active calculations and plotting method calls
        # Assuming active data has been pre-processed similarly to how passive data is handled
        data = read_active_file(active_path)

        # Unpack the data
        (
            frequency,
            start_phi,
            start_theta,
            stop_phi,
            stop_theta,
            inc_phi,
            inc_theta,
            calc_trp,
            theta_angles_deg,
            phi_angles_deg,
            h_power_dBm,
            v_power_dBm,
        ) = (
            data["Frequency"],
            data["Start Phi"],
            data["Start Theta"],
            data["Stop Phi"],
            data["Stop Theta"],
            data["Inc Phi"],
            data["Inc Theta"],
            data["Calculated TRP(dBm)"],
            data["Theta_Angles_Deg"],
            data["Phi_Angles_Deg"],
            data["H_Power_dBm"],
            data["V_Power_dBm"],
        )

        # Calculate active variables
        active_variables = calculate_active_variables(
            start_phi,
            stop_phi,
            start_theta,
            stop_theta,
            inc_phi,
            inc_theta,
            h_power_dBm,
            v_power_dBm,
        )

        # Unpack calculated active variables
        (
            data_points,
            theta_angles_deg,
            phi_angles_deg,
            theta_angles_rad,
            phi_angles_rad,
            total_power_dBm_2d,
            h_power_dBm_2d,
            v_power_dBm_2d,
            phi_angles_deg_plot,
            phi_angles_rad_plot,
            total_power_dBm_2d_plot,
            h_power_dBm_2d_plot,
            v_power_dBm_2d_plot,
            total_power_dBm_min,
            total_power_dBm_nom,
            h_power_dBm_min,
            h_power_dBm_nom,
            v_power_dBm_min,
            v_power_dBm_nom,
            TRP_dBm,
            h_TRP_dBm,
            v_TRP_dBm,
        ) = active_variables

        # Plot and save the 2D and 3D data (instead of displaying)
        print("Saving 2D Active Plots…")
        plot_active_2d_data(
            data_points,
            theta_angles_rad,
            phi_angles_rad,
            total_power_dBm_2d,
            frequency,
            save_path=two_d_data_path,
        )
        print("Saving 3D Active Plots...")
        # For total power
        plot_active_3d_data(
            theta_angles_deg,
            phi_angles_deg,
            total_power_dBm_2d,
            phi_angles_deg_plot,
            total_power_dBm_2d_plot,
            frequency,
            power_type="total",
            interpolate=True,
            axis_mode=axis_mode,
            zmin=zmin,
            zmax=zmax,
            save_path=three_d_data_path,
        )

        # For horizontal polarization (hpol)
        plot_active_3d_data(
            theta_angles_deg,
            phi_angles_deg,
            total_power_dBm_2d,
            phi_angles_deg_plot,
            total_power_dBm_2d_plot,
            frequency,
            power_type="hpol",
            interpolate=True,
            axis_mode=axis_mode,
            zmin=zmin,
            zmax=zmax,
            save_path=three_d_data_path,
        )

        # For vertical polarization (vpol)
        plot_active_3d_data(
            theta_angles_deg,
            phi_angles_deg,
            total_power_dBm_2d,
            phi_angles_deg_plot,
            total_power_dBm_2d_plot,
            frequency,
            power_type="vpol",
            interpolate=True,
            axis_mode=axis_mode,
            zmin=zmin,
            zmax=zmax,
            save_path=three_d_data_path,
        )

    elif scan_type == "passive":
        # After reading & parsing, hpol_data and vpol_data will be lists of dictionaries.
        # Each dictionary will represent a frequency point and will contain:
        #'frequency': The frequency value
        #'cal_factor': Calibration factor for that frequency
        #'data': A list of tuples, where each tuple contains (theta, phi, mag, phase).

        (
            parsed_hpol_data,
            start_phi_h,
            stop_phi_h,
            inc_phi_h,
            start_theta_h,
            stop_theta_h,
            inc_theta_h,
        ) = read_passive_file(hpol_path)
        hpol_data = parsed_hpol_data

        (
            parsed_vpol_data,
            start_phi_v,
            stop_phi_v,
            inc_phi_v,
            start_theta_v,
            stop_theta_v,
            inc_theta_v,
        ) = read_passive_file(vpol_path)
        vpol_data = parsed_vpol_data

        # Check to see if selected files have mismatched frequency or angle data
        angles_match(
            start_phi_h,
            stop_phi_h,
            inc_phi_h,
            start_theta_h,
            stop_theta_h,
            inc_theta_h,
            start_phi_v,
            stop_phi_v,
            inc_phi_v,
            start_theta_v,
            stop_theta_v,
            inc_theta_v,
        )

        # Call Methods to Calculate Required Variables and Set up variables for plotting
        passive_variables = calculate_passive_variables(
            hpol_data,
            vpol_data,
            cable_loss,
            start_phi_h,
            stop_phi_h,
            inc_phi_h,
            start_theta_h,
            stop_theta_h,
            inc_theta_h,
            freq_list,
            selected_frequency,
        )
        theta_angles_deg, phi_angles_deg, v_gain_dB, h_gain_dB, Total_Gain_dB = passive_variables

        # Call Method to Plot Passive Data
        print("Plotting 2D Passive Data...")
        plot_2d_passive_data(
            theta_angles_deg,
            phi_angles_deg,
            v_gain_dB,
            h_gain_dB,
            Total_Gain_dB,
            freq_list,
            selected_frequency,
            datasheet_plots,
            save_path=two_d_data_path,
        )

        print("Plotting 3D Passive Data...")
        for pol in ("total", "hpol", "vpol"):
            plot_passive_3d_component(
                theta_angles_deg,
                phi_angles_deg,
                h_gain_dB if pol == "hpol" else (v_gain_dB if pol == "vpol" else Total_Gain_dB),
                v_gain_dB if pol == "hpol" else (h_gain_dB if pol == "vpol" else Total_Gain_dB),
                Total_Gain_dB,
                freq_list,
                selected_frequency,
                gain_type=pol,
                axis_mode=axis_mode,
                zmin=zmin,
                zmax=zmax,
                save_path=user_selected_frequency_path,
            )
    print(f"Data saved to {project_path}")
