"""
Report Tools for RFlect MCP Server

Generates professional branded DOCX reports with embedded plots, gain tables,
AI-generated summaries/conclusions/captions, and graceful fallback when no
AI provider is configured.
"""

import os
import tempfile
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
from datetime import datetime

from .import_tools import get_loaded_measurements, LoadedMeasurement
from .analysis_tools import (
    _get_analyzer_for_measurement,
    get_gain_statistics,
    analyze_pattern,
    compare_polarizations,
    get_all_analysis,
)


@dataclass
class ReportOptions:
    """Configuration options for report generation."""
    # Content filtering
    frequencies: Optional[List[float]] = None  # None = all frequencies
    polarizations: List[str] = field(default_factory=lambda: ["total"])  # total, hpol, vpol
    measurements: Optional[List[str]] = None  # None = all loaded measurements

    # Plot filtering (key for managing complexity)
    include_2d_plots: bool = True
    include_3d_plots: bool = False  # Default off - they're large/complex
    include_polar_plots: bool = True
    include_cartesian_plots: bool = False
    max_plot_frequencies: int = 5  # Max frequencies for per-freq plots (azimuth cuts)

    # Data filtering
    include_raw_data_tables: bool = False
    include_gain_tables: bool = True
    max_frequencies_in_table: int = 10  # Limit table rows

    # AI content
    ai_executive_summary: bool = True
    ai_section_analysis: bool = True
    ai_recommendations: bool = True
    ai_model: str = "gpt-4o-mini"  # Cost-effective default

    # Output
    output_format: str = "docx"  # docx, pdf (future)
    include_cover_page: bool = True
    include_table_of_contents: bool = True

    # Template
    template_path: Optional[str] = None  # Path to YAML template (None = default)

    # Branding (loaded from config)
    company_name: Optional[str] = None
    logo_path: Optional[str] = None


# ---------------------------------------------------------------------------
# LLM Provider Helper
# ---------------------------------------------------------------------------

def _create_llm_provider(opts: ReportOptions):
    """Create an LLM provider for report generation based on config.

    Returns BaseLLMProvider or None if no API key is configured.
    """
    try:
        from plot_antenna.api_keys import get_api_key
        from plot_antenna.llm_provider import create_provider
        from plot_antenna import config

        ai_provider = getattr(config, "AI_PROVIDER", "openai")

        if ai_provider == "openai":
            api_key = get_api_key("openai")
            if not api_key:
                return None
            model = getattr(config, "AI_OPENAI_MODEL", opts.ai_model)
            return create_provider("openai", api_key=api_key, model=model)
        elif ai_provider == "anthropic":
            api_key = get_api_key("anthropic")
            if not api_key:
                return None
            model = getattr(config, "AI_ANTHROPIC_MODEL", "claude-sonnet-4-20250514")
            return create_provider("anthropic", api_key=api_key, model=model)
        elif ai_provider == "ollama":
            model = getattr(config, "AI_OLLAMA_MODEL", "llama3.1")
            base_url = getattr(config, "AI_OLLAMA_URL", "http://localhost:11434")
            return create_provider("ollama", model=model, base_url=base_url)
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# Plot Generation
# ---------------------------------------------------------------------------

def _generate_plots(measurements: Dict[str, LoadedMeasurement], opts: ReportOptions,
                    plot_dir: str) -> Dict[str, List[str]]:
    """Generate PNG plots for all measurements in headless mode.

    Returns dict mapping measurement_name -> list of image paths.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from plot_antenna.plotting import (
        plot_2d_passive_data,
        plot_passive_3d_component,
        plot_active_2d_data,
        plot_active_3d_data,
    )

    images: Dict[str, List[str]] = {}

    for name, m in measurements.items():
        if opts.measurements is not None and name not in opts.measurements:
            continue

        meas_dir = os.path.join(plot_dir, _safe_filename(name))
        os.makedirs(meas_dir, exist_ok=True)

        meas_images: List[str] = []

        if m.scan_type == "passive":
            meas_images.extend(
                _generate_passive_plots(m, opts, meas_dir,
                                        plot_2d_passive_data, plot_passive_3d_component, plt)
            )
        elif m.scan_type == "active":
            meas_images.extend(
                _generate_active_plots(m, opts, meas_dir,
                                       plot_active_2d_data, plot_active_3d_data, plt)
            )

        images[name] = meas_images

    return images


def _generate_passive_plots(m: LoadedMeasurement, opts: ReportOptions,
                            meas_dir: str, plot_2d_fn, plot_3d_fn, plt) -> List[str]:
    """Generate passive measurement plots. Returns list of image paths."""
    paths: List[str] = []
    data = m.data
    theta = data.get("theta")
    phi = data.get("phi")
    v_gain = data.get("v_gain")
    h_gain = data.get("h_gain")
    total_gain = data.get("total_gain")

    if theta is None or phi is None or total_gain is None:
        return paths

    freqs = _filter_frequencies(m.frequencies, opts)
    band_info = _detect_rf_band(freqs)
    plot_freqs = _select_representative_frequencies(freqs, opts.max_plot_frequencies,
                                                    band_info=band_info)

    # 2D plots: summary plots once (at first freq), azimuth cuts at representative freqs
    if opts.include_2d_plots:
        save_dir_2d = os.path.join(meas_dir, "2d")
        os.makedirs(save_dir_2d, exist_ok=True)
        for freq in plot_freqs:
            try:
                plot_2d_fn(
                    theta, phi, v_gain, h_gain, total_gain,
                    m.frequencies, freq,
                    datasheet_plots=False,
                    save_path=save_dir_2d,
                )
                plt.close("all")
            except Exception:
                plt.close("all")
        # Collect all generated PNGs
        for f in sorted(os.listdir(save_dir_2d)):
            full = os.path.join(save_dir_2d, f)
            if full.endswith(".png"):
                paths.append(full)

    # 3D plots: only at center frequency (filenames don't include freq,
    # so generating for multiple freqs just overwrites the same files)
    if opts.include_3d_plots and freqs:
        center_freq = freqs[len(freqs) // 2]
        save_dir_3d = os.path.join(meas_dir, "3d")
        os.makedirs(save_dir_3d, exist_ok=True)
        for gain_type in ("total", "hpol", "vpol"):
            if gain_type != "total" and gain_type not in opts.polarizations:
                continue
            try:
                plot_3d_fn(
                    theta, phi, h_gain, v_gain, total_gain,
                    m.frequencies, center_freq,
                    gain_type=gain_type,
                    save_path=save_dir_3d,
                )
                plt.close("all")
            except Exception:
                plt.close("all")
        for f in sorted(os.listdir(save_dir_3d)):
            full = os.path.join(save_dir_3d, f)
            if full.endswith(".png"):
                paths.append(full)

    return paths


def _generate_active_plots(m: LoadedMeasurement, opts: ReportOptions,
                           meas_dir: str, plot_2d_fn, plot_3d_fn, plt) -> List[str]:
    """Generate active measurement plots. Returns list of image paths."""
    paths: List[str] = []
    data = m.data

    # Need the 2D arrays stored by import_active_processed
    data_points = data.get("data_points")
    theta_rad = data.get("theta_rad")
    phi_rad_plot = data.get("phi_rad_plot")
    total_power_2d_plot = data.get("total_power_2d_plot")

    # For 3D
    theta_deg = data.get("theta")
    phi_deg = data.get("phi")
    total_power_2d = data.get("total_power_2d")
    phi_deg_plot = data.get("phi_deg_plot")

    freq = m.frequencies[0] if m.frequencies else 0

    if opts.include_2d_plots and all(v is not None for v in
                                     [data_points, theta_rad, phi_rad_plot, total_power_2d_plot]):
        save_dir_2d = os.path.join(meas_dir, "2d")
        os.makedirs(save_dir_2d, exist_ok=True)
        try:
            plot_2d_fn(data_points, theta_rad, phi_rad_plot,
                       total_power_2d_plot, freq, save_path=save_dir_2d)
            plt.close("all")
            for f in sorted(os.listdir(save_dir_2d)):
                full = os.path.join(save_dir_2d, f)
                if full.endswith(".png") and full not in paths:
                    paths.append(full)
        except Exception:
            plt.close("all")

    if opts.include_3d_plots and all(v is not None for v in
                                     [theta_deg, phi_deg, total_power_2d,
                                      phi_deg_plot, total_power_2d_plot]):
        save_dir_3d = os.path.join(meas_dir, "3d")
        os.makedirs(save_dir_3d, exist_ok=True)
        try:
            plot_3d_fn(theta_deg, phi_deg, total_power_2d,
                       phi_deg_plot, total_power_2d_plot, freq,
                       save_path=save_dir_3d)
            plt.close("all")
            for f in sorted(os.listdir(save_dir_3d)):
                full = os.path.join(save_dir_3d, f)
                if full.endswith(".png") and full not in paths:
                    paths.append(full)
        except Exception:
            plt.close("all")

    return paths


def _filter_frequencies(all_freqs: List[float], opts: ReportOptions) -> List[float]:
    """Return frequencies filtered by opts, or all if no filter."""
    if opts.frequencies is None:
        return list(all_freqs)
    return [f for f in all_freqs if f in opts.frequencies]


# ---------------------------------------------------------------------------
# RF Band Detection
# ---------------------------------------------------------------------------

# Band definitions: (name, standard, min_mhz, max_mhz, key_frequencies, channels)
_RF_BANDS = [
    {
        "name": "BLE",
        "standard": "Bluetooth 5.x",
        "range": (2400, 2483.5),
        "key_frequencies": [2402, 2426, 2440, 2454, 2480],
        "channels": {
            "CH37 (adv)": 2402, "CH38 (adv)": 2426, "CH18 (center)": 2440,
            "CH25": 2454, "CH39 (adv)": 2480,
        },
    },
    {
        "name": "WiFi 2.4 GHz",
        "standard": "IEEE 802.11",
        "range": (2400, 2500),
        "key_frequencies": [2412, 2437, 2462],
        "channels": {"CH1": 2412, "CH6": 2437, "CH11": 2462},
    },
    {
        "name": "WiFi 5 GHz",
        "standard": "IEEE 802.11ac/ax",
        "range": (5150, 5850),
        "key_frequencies": [5180, 5500, 5745, 5825],
        "channels": {"CH36": 5180, "CH100": 5500, "CH149": 5745, "CH165": 5825},
    },
    {
        "name": "WiFi 6E",
        "standard": "IEEE 802.11ax",
        "range": (5925, 7125),
        "key_frequencies": [5955, 6415, 6875, 7115],
        "channels": {},
    },
    {
        "name": "LoRa EU868",
        "standard": "EU868",
        "range": (863, 870),
        "key_frequencies": [868.1, 868.3, 868.5],
        "channels": {},
    },
    {
        "name": "LoRa US915",
        "standard": "US915",
        "range": (902, 928),
        "key_frequencies": [903, 915, 927],
        "channels": {},
    },
    {
        "name": "GPS L1",
        "standard": "GPS",
        "range": (1565.42, 1585.42),
        "key_frequencies": [1575.42],
        "channels": {"L1": 1575.42},
    },
    {
        "name": "LTE Band 7",
        "standard": "3GPP",
        "range": (2500, 2690),
        "key_frequencies": [2535, 2595, 2655],
        "channels": {},
    },
    {
        "name": "Sub-GHz IoT",
        "standard": "ISM",
        "range": (315, 435),
        "key_frequencies": [315, 433.92],
        "channels": {},
    },
    {
        "name": "UWB",
        "standard": "IEEE 802.15.4a/z",
        "range": (3100, 10600),
        "key_frequencies": [3494.4, 3993.6, 4492.8, 6489.6, 7987.2],
        "channels": {
            "CH1": 3494.4, "CH2": 3993.6, "CH3": 4492.8,
            "CH5": 6489.6, "CH9": 7987.2,
        },
    },
]


def _detect_rf_band(frequencies: List[float]) -> Optional[Dict]:
    """Detect which RF band a set of frequencies belongs to.

    Checks if >60% of frequencies fall within a known band range.
    Returns the band with the best coverage, or None if no match.

    Returns dict with: name, standard, full_range, key_frequencies, channels.
    """
    if not frequencies:
        return None

    best_band = None
    best_coverage = 0.0

    for band in _RF_BANDS:
        lo, hi = band["range"]
        in_band = sum(1 for f in frequencies if lo <= f <= hi)
        coverage = in_band / len(frequencies)
        if coverage > best_coverage:
            best_coverage = coverage
            best_band = band

    if best_coverage < 0.6:
        return None

    return {
        "name": best_band["name"],
        "standard": best_band["standard"],
        "full_range": best_band["range"],
        "key_frequencies": best_band["key_frequencies"],
        "channels": best_band["channels"],
    }


def _snap_to_nearest(target: float, available: List[float],
                     tolerance_mhz: float = 5.0) -> Optional[float]:
    """Find the closest frequency in available list to target, within tolerance."""
    if not available:
        return None
    closest = min(available, key=lambda f: abs(f - target))
    if abs(closest - target) <= tolerance_mhz:
        return closest
    return None


def _channel_label(freq: float, band_info: Optional[Dict]) -> Optional[str]:
    """Get channel label for a frequency, e.g. 'BLE CH39 (adv)' for 2480 MHz."""
    if band_info is None:
        return None
    for label, ch_freq in band_info.get("channels", {}).items():
        if abs(freq - ch_freq) < 1.0:
            return f"{band_info['name']} {label}"
    return None


def _select_representative_frequencies(freqs: List[float], max_count: int,
                                       band_info: Optional[Dict] = None) -> List[float]:
    """Select representative frequencies, preferring key band channels.

    If band_info is provided, starts with key channel frequencies that exist
    in the dataset (snapped to nearest available), always includes first/last,
    then fills remaining slots with evenly-spaced picks.

    Falls back to pure even-spacing when no band is detected.
    """
    if len(freqs) <= max_count:
        return list(freqs)
    if max_count <= 0:
        return []
    if max_count == 1:
        return [freqs[len(freqs) // 2]]

    if band_info is not None:
        # Start with key channel frequencies snapped to available data
        selected = []
        for key_freq in band_info.get("key_frequencies", []):
            snapped = _snap_to_nearest(key_freq, freqs)
            if snapped is not None and snapped not in selected:
                selected.append(snapped)

        # Always include first and last
        if freqs[0] not in selected:
            selected.insert(0, freqs[0])
        if freqs[-1] not in selected:
            selected.append(freqs[-1])

        # If we already have enough, trim to max_count keeping edges + key freqs
        if len(selected) > max_count:
            # Keep first, last, and as many key freqs as fit
            edges = {freqs[0], freqs[-1]}
            key_only = [f for f in selected if f not in edges]
            selected = [freqs[0]] + key_only[:max_count - 2] + [freqs[-1]]

        # Fill remaining slots with evenly-spaced picks
        if len(selected) < max_count:
            remaining = [f for f in freqs if f not in selected]
            slots = max_count - len(selected)
            if remaining and slots > 0:
                step = max(1, len(remaining) // (slots + 1))
                for i in range(1, slots + 1):
                    idx = min(i * step, len(remaining) - 1)
                    if remaining[idx] not in selected:
                        selected.append(remaining[idx])

        # Sort by frequency
        selected.sort()
        return selected[:max_count]

    # Fallback: evenly-spaced algorithm
    indices = [0]
    for i in range(1, max_count - 1):
        idx = int(round(i * (len(freqs) - 1) / (max_count - 1)))
        if idx not in indices:
            indices.append(idx)
    indices.append(len(freqs) - 1)
    # Deduplicate while preserving order
    seen = set()
    result = []
    for idx in indices:
        if idx not in seen:
            seen.add(idx)
            result.append(freqs[idx])
    return result


def _safe_filename(name: str) -> str:
    """Convert a measurement name to a safe directory name."""
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()


def _pretty_caption(filename: str) -> str:
    """Convert a plot filename to a human-readable caption.

    'Azimuth_Cuts_2450.0MHz.png' -> 'Azimuth Cuts at 2450.0 MHz'
    '3D_total_1of2.png' -> '3D Total Pattern (View 1)'
    'efficiency_db_2300-2600MHz.png' -> 'Efficiency (dB) 2300-2600 MHz'
    """
    import re
    name = filename.rsplit(".", 1)[0]  # strip extension

    # 3D plot views
    m = re.match(r"3D_(\w+)_(\d+)of(\d+)", name)
    if m:
        label = m.group(1).replace("_", " ").title()
        return f"3D {label} Pattern (View {m.group(2)} of {m.group(3)})"

    # Active 3D TRP plots
    m = re.match(r"3D_TRP_(\w+)_(.+?)MHz_(\d+)of(\d+)", name)
    if m:
        label = m.group(1).replace("_", " ").title()
        return f"3D TRP {label} at {m.group(2)} MHz (View {m.group(3)} of {m.group(4)})"

    # Efficiency plots
    m = re.match(r"efficiency_(db|percent)_(.+?)MHz", name)
    if m:
        unit = "dB" if m.group(1) == "db" else "%"
        return f"Radiated Efficiency ({unit}) {m.group(2)} MHz"

    # Total gain vs freq
    m = re.match(r"total_gain_(.+?)MHz", name)
    if m:
        return f"Peak Gain vs Frequency {m.group(1)} MHz"

    # Azimuth cuts
    m = re.match(r"Azimuth_Cuts_(.+?)MHz", name)
    if m:
        return f"Azimuth Pattern Cuts at {m.group(1)} MHz"

    # Active 2D azimuth
    m = re.match(r"2D_Azimuth_Cuts_(.+?)_MHz", name)
    if m:
        return f"Azimuth Power Cuts at {m.group(1)} MHz"

    # Elevation power pattern (active)
    m = re.match(r"Elevation_Power_Pattern_at_Phi___(\d+)_(\d+)deg,_Freq___(.+?)MHz", name)
    if m:
        return f"Elevation Power Pattern at Phi={m.group(1)}/{m.group(2)} deg, {m.group(3)} MHz"

    # Total power in theta plane (active)
    m = re.match(r"Total_Power_in_Theta___(\d+)deg_Plane,_Freq___(.+?)MHz", name)
    if m:
        return f"Total Power in Theta={m.group(1)} deg Plane, {m.group(2)} MHz"

    # Generic: replace underscores, strip trailing freq
    return name.replace("_", " ")


def _sort_images_by_frequency(img_paths: List[str]) -> List[str]:
    """Sort image paths so frequency-specific plots are in frequency order."""
    import re

    def sort_key(path: str):
        basename = os.path.basename(path)
        # Extract frequency number from filename
        m = re.search(r"(\d+\.?\d*)(?:MHz|_MHz)", basename)
        freq = float(m.group(1)) if m else 0
        # Group: summary plots first (no per-freq), then by frequency
        is_per_freq = bool(re.search(r"Azimuth|2D_Azimuth|3D_", basename))
        return (0 if not is_per_freq else 1, freq, basename)

    return sorted(img_paths, key=sort_key)


# ---------------------------------------------------------------------------
# AI Text Generation
# ---------------------------------------------------------------------------

def _generate_ai_text(provider, prompt: str, data: Dict, opts: ReportOptions,
                      max_tokens: int = 500) -> Optional[str]:
    """Generate AI text using an LLM provider. Returns None on failure."""
    if provider is None:
        return None
    try:
        from plot_antenna.llm_provider import LLMMessage

        all_analysis = []
        for freq in data["frequencies"][:3]:
            all_analysis.append(get_all_analysis(freq))

        full_prompt = f"{prompt}\n\nMeasurement Data:\n" + "\n".join(all_analysis)
        response = provider.chat(
            [LLMMessage(role="user", content=full_prompt)], max_tokens=max_tokens
        )
        return response.content or None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Branded DOCX Builder
# ---------------------------------------------------------------------------

def _fmt(val, fmt=".2f", suffix=""):
    """Format a value for table display, handling None gracefully."""
    if val is None:
        return "N/A"
    try:
        return f"{float(val):{fmt}}{suffix}"
    except (ValueError, TypeError):
        return str(val)


# ---------------------------------------------------------------------------
# Classification Helpers
# ---------------------------------------------------------------------------

def _classify_gain_quality(peak_gain_dBi, freq_mhz=None):
    """Classify antenna gain quality against engineering benchmarks.

    Returns (rating, description) tuple.
    """
    if peak_gain_dBi is None:
        return "unknown", "insufficient data"
    g = float(peak_gain_dBi)
    if g < 0:
        return "poor", "below isotropic (0 dBi)"
    elif g < 2:
        return "low", "chip antenna range (0-2 dBi)"
    elif g < 5:
        return "moderate", "PCB trace / PIFA range (2-5 dBi)"
    elif g < 8:
        return "good", "patch antenna range (5-8 dBi)"
    elif g < 12:
        return "high", "directional antenna range (8-12 dBi)"
    else:
        return "very high", "high-gain directional (>12 dBi)"


def _classify_efficiency(efficiency_pct):
    """Classify antenna efficiency against standard benchmarks."""
    if efficiency_pct is None:
        return "unknown"
    e = float(efficiency_pct)
    if e > 90:
        return "excellent (>90%)"
    elif e > 70:
        return "good (70-90%)"
    elif e > 50:
        return "fair (50-70%)"
    else:
        return "poor (<50%)"


def _get_test_configuration(m: LoadedMeasurement) -> Dict:
    """Extract test configuration details from a measurement."""
    import numpy as np
    config_info = {
        "scan_type": m.scan_type,
        "num_frequencies": len(m.frequencies),
    }
    if len(m.frequencies) > 1:
        config_info["freq_range"] = f"{min(m.frequencies):.0f} - {max(m.frequencies):.0f} MHz"
        config_info["freq_step"] = f"{m.frequencies[1] - m.frequencies[0]:.1f} MHz"
    elif m.frequencies:
        config_info["freq_range"] = f"{m.frequencies[0]:.0f} MHz"

    data = m.data
    theta = data.get("theta")
    phi = data.get("phi")
    if theta is not None and hasattr(theta, '__len__'):
        unique_theta = np.unique(theta)
        config_info["theta_range"] = f"{float(np.min(theta)):.0f} - {float(np.max(theta)):.0f} deg"
        config_info["theta_points"] = len(unique_theta)
        if len(unique_theta) > 1:
            config_info["theta_step"] = f"{float(unique_theta[1] - unique_theta[0]):.1f} deg"
    if phi is not None and hasattr(phi, '__len__'):
        unique_phi = np.unique(phi)
        config_info["phi_range"] = f"{float(np.min(phi)):.0f} - {float(np.max(phi)):.0f} deg"
        config_info["phi_points"] = len(unique_phi)
        if len(unique_phi) > 1:
            config_info["phi_step"] = f"{float(unique_phi[1] - unique_phi[0]):.1f} deg"

    if m.scan_type == "active":
        if "TRP_dBm" in data:
            config_info["TRP_dBm"] = float(data["TRP_dBm"])
        if "data_points" in data:
            config_info["spatial_points"] = data["data_points"]

    return config_info


# ---------------------------------------------------------------------------
# Data-Driven Text Generators
# ---------------------------------------------------------------------------

def _build_executive_summary(measurements: Dict[str, LoadedMeasurement],
                              opts: ReportOptions,
                              band_info_map: Optional[Dict[str, Optional[Dict]]] = None
                              ) -> List[str]:
    """Build a data-driven executive summary (3 paragraphs, no AI needed)."""
    paragraphs = []

    # Classify measurements
    passive_names = [n for n, m in measurements.items()
                     if m.scan_type == "passive"
                     and (opts.measurements is None or n in opts.measurements)]
    active_names = [n for n, m in measurements.items()
                    if m.scan_type == "active"
                    and (opts.measurements is None or n in opts.measurements)]

    all_freqs = []
    for n in passive_names + active_names:
        all_freqs.extend(measurements[n].frequencies)
    all_freqs = sorted(set(all_freqs))

    # Detect overall band (use first detected band across measurements)
    overall_band = None
    if band_info_map:
        for name in passive_names + active_names:
            bi = band_info_map.get(name)
            if bi is not None:
                overall_band = bi
                break
    if overall_band is None:
        overall_band = _detect_rf_band(all_freqs)

    # Paragraph 1: Scope
    parts = []
    if passive_names:
        parts.append(f"{len(passive_names)} passive gain measurement(s)")
    if active_names:
        parts.append(f"{len(active_names)} active TRP measurement(s)")
    scope = f"This report documents {' and '.join(parts)}"
    if len(all_freqs) > 1:
        scope += f" spanning {min(all_freqs):.0f} to {max(all_freqs):.0f} MHz"
        scope += f" ({len(all_freqs)} frequency points"
        if overall_band:
            scope += f", {overall_band['name']} band"
        scope += ")"
    elif all_freqs:
        scope += f" at {all_freqs[0]:.0f} MHz"
    scope += "."
    paragraphs.append(scope)

    # Paragraph 2: Performance highlights
    highlights = []
    for name in passive_names:
        m = measurements[name]
        analyzer, _, err = _get_analyzer_for_measurement(name)
        if analyzer is None:
            continue
        overall = analyzer.analyze_all_frequencies()
        res_freq = overall.get("resonance_frequency_MHz")
        peak_gain = overall.get("peak_gain_at_resonance_dBi")
        bw = overall.get("bandwidth_3dB_MHz")
        variation = overall.get("gain_variation_dB")

        line = f"{name}: peak gain {_fmt(peak_gain)} dBi"
        if res_freq is not None:
            line += f" at {_fmt(res_freq, '.0f')} MHz"
        rating, desc = _classify_gain_quality(peak_gain)
        line += f" ({desc})"
        if bw is not None:
            line += f", 3 dB bandwidth {_fmt(bw, '.0f')} MHz"
        if variation is not None and variation > 3.0:
            line += f", gain variation {_fmt(variation)} dB (significant)"

        # Efficiency at resonance
        if res_freq is not None:
            pattern = analyzer.analyze_pattern(frequency=res_freq)
            eff = pattern.get("estimated_efficiency_pct")
            if eff is not None:
                line += f", estimated efficiency {_fmt(eff, '.0f')}% ({_classify_efficiency(eff)})"

        highlights.append(line)

    for name in active_names:
        m = measurements[name]
        data = m.data
        trp = data.get("TRP_dBm")
        h_trp = data.get("h_TRP_dBm")
        v_trp = data.get("v_TRP_dBm")
        freq = m.frequencies[0] if m.frequencies else 0

        line = f"{name}: TRP = {_fmt(trp)} dBm at {_fmt(freq, '.0f')} MHz"
        if h_trp is not None and v_trp is not None:
            balance = abs(float(h_trp) - float(v_trp))
            line += f" (H/V balance: {_fmt(balance, '.1f')} dB)"
        highlights.append(line)

    if highlights:
        paragraphs.append("Performance highlights: " + "; ".join(highlights) + ".")

    # Paragraph 3: Observations / flags
    observations = []
    for name in passive_names:
        analyzer, _, _ = _get_analyzer_for_measurement(name)
        if analyzer is None:
            continue
        overall = analyzer.analyze_all_frequencies()
        peak = overall.get("peak_gain_at_resonance_dBi")
        variation = overall.get("gain_variation_dB")
        if peak is not None and peak < 0:
            observations.append(f"{name} exhibits sub-isotropic peak gain ({_fmt(peak)} dBi)")
        if variation is not None and variation > 6:
            observations.append(
                f"{name} shows high gain variation ({_fmt(variation)} dB) across the band")
    if observations:
        paragraphs.append("Observations: " + ". ".join(observations) + ".")

    return paragraphs


def _build_pattern_prose(analyzer, freq, measurement_name,
                         band_info: Optional[Dict] = None) -> str:
    """Build a prose paragraph describing pattern characteristics at a frequency."""
    pattern = analyzer.analyze_pattern(frequency=freq)
    stats = analyzer.get_gain_statistics(frequency=freq)

    peak = pattern.get("peak_gain_dBi")
    pattern_type = pattern.get("pattern_type", "unknown")
    hpbw_e = pattern.get("hpbw_e_plane")
    hpbw_h = pattern.get("hpbw_h_plane")
    fb = pattern.get("front_to_back_dB")
    eff = pattern.get("estimated_efficiency_pct")
    beam_theta = pattern.get("main_beam_theta")
    beam_phi = pattern.get("main_beam_phi")

    actual_freq = pattern.get("frequency", freq)

    # Build frequency label with optional channel annotation
    freq_label = f"{_fmt(actual_freq, '.0f')} MHz"
    ch_label = _channel_label(actual_freq, band_info)
    if ch_label:
        freq_label += f" ({ch_label})"

    parts = [f"At {freq_label}, {measurement_name} exhibits "
             f"a {pattern_type} radiation pattern with peak gain of {_fmt(peak)} dBi"]

    rating, desc = _classify_gain_quality(peak)
    parts[0] += f" ({desc})."

    beam_parts = []
    if hpbw_e is not None:
        beam_parts.append(f"E-plane HPBW: {_fmt(hpbw_e, '.1f')}\u00b0")
    if hpbw_h is not None:
        beam_parts.append(f"H-plane HPBW: {_fmt(hpbw_h, '.1f')}\u00b0")
    if fb is not None:
        beam_parts.append(f"front-to-back ratio: {_fmt(fb, '.1f')} dB")
    if beam_parts:
        parts.append("Beam characteristics: " + ", ".join(beam_parts) + ".")

    if beam_theta is not None and beam_phi is not None:
        parts.append(
            f"Main beam direction: \u03b8={_fmt(beam_theta, '.0f')}\u00b0, "
            f"\u03c6={_fmt(beam_phi, '.0f')}\u00b0.")

    if eff is not None:
        parts.append(
            f"Estimated radiation efficiency: {_fmt(eff, '.0f')}% "
            f"({_classify_efficiency(eff)}).")

    num_nulls = pattern.get("num_nulls", 0)
    if num_nulls > 0:
        deepest = pattern.get("deepest_null_dB")
        null_text = f"{num_nulls} null(s) detected"
        if deepest is not None:
            null_text += f", deepest at {_fmt(deepest)} dBi"
        parts.append(null_text + ".")

    return " ".join(parts)


def _build_data_driven_conclusions(measurements: Dict[str, LoadedMeasurement],
                                    opts: ReportOptions,
                                    band_info_map: Optional[Dict[str, Optional[Dict]]] = None
                                    ) -> List[str]:
    """Build data-driven conclusion bullets referencing actual values."""
    bullets = []

    passive_names = [n for n, m in measurements.items()
                     if m.scan_type == "passive"
                     and (opts.measurements is None or n in opts.measurements)]
    active_names = [n for n, m in measurements.items()
                    if m.scan_type == "active"
                    and (opts.measurements is None or n in opts.measurements)]

    for name in passive_names:
        analyzer, _, err = _get_analyzer_for_measurement(name)
        if analyzer is None:
            continue
        overall = analyzer.analyze_all_frequencies()
        peak = overall.get("peak_gain_at_resonance_dBi")
        res_freq = overall.get("resonance_frequency_MHz")
        bw = overall.get("bandwidth_3dB_MHz")
        variation = overall.get("gain_variation_dB")

        # Gain assessment
        if peak is not None:
            rating, desc = _classify_gain_quality(peak)
            bullets.append(
                f"{name}: Peak gain of {_fmt(peak)} dBi at {_fmt(res_freq, '.0f')} MHz "
                f"places this antenna in the {desc} category")

        # Bandwidth + band coverage
        if bw is not None:
            bw_text = (f"{name}: 3 dB bandwidth of {_fmt(bw, '.0f')} MHz "
                       f"({_fmt(res_freq, '.0f')} MHz center)")
            meas_band = (band_info_map or {}).get(name)
            if meas_band is not None:
                band_lo, band_hi = meas_band["full_range"]
                band_width = band_hi - band_lo
                if bw is not None and float(bw) >= band_width:
                    bw_text += (f" covers the full {meas_band['name']} band "
                                f"({band_lo:.0f}\u2013{band_hi:.0f} MHz)")
                else:
                    coverage_pct = min(100, float(bw) / band_width * 100) if band_width > 0 else 0
                    bw_text += (f" covers {coverage_pct:.0f}% of the {meas_band['name']} band "
                                f"({band_lo:.0f}\u2013{band_hi:.0f} MHz)")
            bullets.append(bw_text)

        # Gain stability
        if variation is not None:
            if variation > 6:
                bullets.append(
                    f"{name}: Gain variation of {_fmt(variation)} dB across band is "
                    f"significant; design optimization may improve flatness")
            elif variation > 3:
                bullets.append(
                    f"{name}: Gain variation of {_fmt(variation)} dB across band is "
                    f"moderate; acceptable for most applications")

        # Efficiency at resonance
        if res_freq is not None:
            pattern = analyzer.analyze_pattern(frequency=res_freq)
            eff = pattern.get("estimated_efficiency_pct")
            if eff is not None:
                eff_class = _classify_efficiency(eff)
                if eff < 50:
                    bullets.append(
                        f"{name}: Estimated efficiency of {_fmt(eff, '.0f')}% ({eff_class}) "
                        f"suggests significant ohmic or mismatch losses")
                elif eff < 70:
                    bullets.append(
                        f"{name}: Estimated efficiency of {_fmt(eff, '.0f')}% ({eff_class}); "
                        f"impedance matching improvements may yield gains")

        # Polarization balance
        if analyzer.frequencies:
            mid_freq = analyzer.frequencies[len(analyzer.frequencies) // 2]
            pol = analyzer.compare_polarizations(frequency=mid_freq)
            balance = pol.get("polarization_balance_dB")
            if balance is not None and abs(balance) > 6:
                dominant = pol.get("dominant_pol", "?")
                bullets.append(
                    f"{name}: Polarization imbalance of {_fmt(abs(balance), '.1f')} dB "
                    f"({dominant}-pol dominant); consider design adjustments for "
                    f"balanced polarization if required")

    for name in active_names:
        m = measurements[name]
        data = m.data
        trp = data.get("TRP_dBm")
        h_trp = data.get("h_TRP_dBm")
        v_trp = data.get("v_TRP_dBm")

        if trp is not None:
            bullets.append(
                f"{name}: Total Radiated Power of {_fmt(trp)} dBm")
        if h_trp is not None and v_trp is not None:
            balance = abs(float(h_trp) - float(v_trp))
            if balance > 3:
                bullets.append(
                    f"{name}: H/V TRP imbalance of {_fmt(balance, '.1f')} dB "
                    f"indicates polarization skew in the radiated power")

    if not bullets:
        bullets.append("Insufficient data to draw specific conclusions")

    return bullets


# ---------------------------------------------------------------------------
# Consolidated Table Builders
# ---------------------------------------------------------------------------

def _style_header_row(table, brand_dark):
    """Apply branded styling to the header row of a table."""
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = brand_dark


def _build_consolidated_performance_table(doc, analyzer, name, rep_freqs,
                                           brand_dark, scan_type):
    """Build ONE consolidated performance table per measurement.

    Passive: Freq | Peak Gain | Avg Gain | Eff% | HPBW-E | HPBW-H | F/B | Pattern
    Active:  Freq | Peak Power | Avg Power | TRP | H-TRP | V-TRP
    """
    from docx.shared import Pt

    if scan_type == "passive":
        headers = ["Freq (MHz)", "Peak (dBi)", "Avg (dBi)", "Eff (%)",
                    "HPBW-E (\u00b0)", "HPBW-H (\u00b0)", "F/B (dB)", "Pattern"]
        table = doc.add_table(rows=1 + len(rep_freqs), cols=len(headers))
        table.style = "Light Shading Accent 1"

        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h
        _style_header_row(table, brand_dark)

        for i, freq in enumerate(rep_freqs):
            stats = analyzer.get_gain_statistics(frequency=freq)
            pattern = analyzer.analyze_pattern(frequency=freq)
            row = table.rows[i + 1]
            row.cells[0].text = _fmt(freq, ".0f")
            row.cells[1].text = _fmt(stats.get("max_gain_dBi"))
            row.cells[2].text = _fmt(stats.get("avg_gain_dBi"))
            row.cells[3].text = _fmt(pattern.get("estimated_efficiency_pct"), ".0f")
            row.cells[4].text = _fmt(pattern.get("hpbw_e_plane"), ".1f")
            row.cells[5].text = _fmt(pattern.get("hpbw_h_plane"), ".1f")
            row.cells[6].text = _fmt(pattern.get("front_to_back_dB"), ".1f")
            row.cells[7].text = str(pattern.get("pattern_type", "N/A"))

    else:  # active
        headers = ["Freq (MHz)", "Peak (dBm)", "Avg (dBm)", "TRP (dBm)",
                    "H-TRP (dBm)", "V-TRP (dBm)"]
        table = doc.add_table(rows=1 + len(rep_freqs), cols=len(headers))
        table.style = "Light Shading Accent 1"

        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h
        _style_header_row(table, brand_dark)

        for i, freq in enumerate(rep_freqs):
            stats = analyzer.get_gain_statistics(frequency=freq)
            row = table.rows[i + 1]
            row.cells[0].text = _fmt(freq, ".0f")
            row.cells[1].text = _fmt(stats.get("max_power_dBm"))
            row.cells[2].text = _fmt(stats.get("avg_power_dBm"))
            row.cells[3].text = _fmt(stats.get("TRP_dBm"))
            row.cells[4].text = _fmt(stats.get("h_TRP_dBm"))
            row.cells[5].text = _fmt(stats.get("v_TRP_dBm"))

    # Summary line under table
    if scan_type == "passive":
        overall = analyzer.analyze_all_frequencies()
        parts = []
        if overall.get("resonance_frequency_MHz"):
            parts.append(
                f"Resonance: {_fmt(overall['resonance_frequency_MHz'], '.0f')} MHz")
            parts.append(
                f"Peak at resonance: {_fmt(overall.get('peak_gain_at_resonance_dBi'))} dBi")
        if overall.get("gain_variation_dB") is not None:
            parts.append(
                f"Gain variation: {_fmt(overall['gain_variation_dB'])} dB")
        if overall.get("bandwidth_3dB_MHz") is not None:
            parts.append(
                f"3 dB BW: {_fmt(overall['bandwidth_3dB_MHz'], '.0f')} MHz")
        if parts:
            summary_para = doc.add_paragraph(" | ".join(parts))
            summary_para.paragraph_format.space_before = Pt(4)
            for run in summary_para.runs:
                run.font.size = Pt(9)
                run.italic = True

    doc.add_paragraph()


def _build_comparison_table(doc, analyzer, rep_freqs, brand_dark, heading_fn):
    """Build a multi-frequency comparison table at representative frequencies only."""
    from docx.shared import Pt

    if len(rep_freqs) < 2:
        return

    heading_fn(doc, "Multi-Frequency Comparison", level=2)

    headers = [
        "Freq (MHz)", "Peak Gain (dBi)", "Pattern Type",
        "HPBW-E (\u00b0)", "HPBW-H (\u00b0)", "F/B (dB)",
    ]
    table = doc.add_table(rows=1 + len(rep_freqs), cols=len(headers))
    table.style = "Light Shading Accent 1"

    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    _style_header_row(table, brand_dark)

    for i, freq in enumerate(rep_freqs):
        pattern = analyzer.analyze_pattern(frequency=freq)
        row = table.rows[i + 1]
        row.cells[0].text = _fmt(freq, ".1f")
        row.cells[1].text = _fmt(pattern.get("peak_gain_dBi"))
        row.cells[2].text = str(pattern.get("pattern_type", "N/A"))
        row.cells[3].text = _fmt(pattern.get("hpbw_e_plane"), ".1f")
        row.cells[4].text = _fmt(pattern.get("hpbw_h_plane"), ".1f")
        row.cells[5].text = _fmt(pattern.get("front_to_back_dB"), ".1f")

    overall = analyzer.analyze_all_frequencies()
    if overall.get("resonance_frequency_MHz"):
        parts = [
            f"Resonance: {_fmt(overall['resonance_frequency_MHz'], '.1f')} MHz",
            f"Peak at resonance: {_fmt(overall.get('peak_gain_at_resonance_dBi'))} dBi",
            f"Gain variation: {_fmt(overall.get('gain_variation_dB'))} dB",
        ]
        if overall.get("bandwidth_3dB_MHz") is not None:
            parts.append(
                f"3 dB BW: {_fmt(overall['bandwidth_3dB_MHz'], '.1f')} MHz")
        summary_para = doc.add_paragraph(" | ".join(parts))
        summary_para.paragraph_format.space_before = Pt(6)
        for run in summary_para.runs:
            run.font.size = Pt(9)
            run.italic = True

    doc.add_paragraph()


def _build_polarization_table(doc, analyzer, rep_freqs, brand_dark):
    """Build ONE consolidated polarization table."""
    headers = ["Freq (MHz)", "HPOL Peak (dBi)", "VPOL Peak (dBi)",
               "Avg XPD (dB)", "Dominant", "Balance (dB)"]
    table = doc.add_table(rows=1 + len(rep_freqs), cols=len(headers))
    table.style = "Light Shading Accent 1"

    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    _style_header_row(table, brand_dark)

    for i, freq in enumerate(rep_freqs):
        pol = analyzer.compare_polarizations(frequency=freq)
        row = table.rows[i + 1]
        row.cells[0].text = _fmt(freq, ".0f")
        row.cells[1].text = _fmt(pol.get("max_hpol_gain_dBi"))
        row.cells[2].text = _fmt(pol.get("max_vpol_gain_dBi"))
        row.cells[3].text = _fmt(pol.get("avg_xpd_dB"), ".1f")
        row.cells[4].text = str(pol.get("dominant_pol", "N/A"))
        row.cells[5].text = _fmt(pol.get("polarization_balance_dB"), ".1f")

    doc.add_paragraph()


def _build_trp_section(doc, m: LoadedMeasurement, brand_dark, heading_fn):
    """Build a dedicated TRP analysis section for active measurements."""
    data = m.data
    trp = data.get("TRP_dBm")
    h_trp = data.get("h_TRP_dBm")
    v_trp = data.get("v_TRP_dBm")
    freq = m.frequencies[0] if m.frequencies else 0

    if trp is None:
        return

    # TRP summary table
    rows = [
        ("Parameter", "Value"),
        ("Frequency", f"{_fmt(freq, '.0f')} MHz"),
        ("Total TRP", f"{_fmt(trp)} dBm"),
    ]
    if h_trp is not None:
        rows.append(("H-pol TRP", f"{_fmt(h_trp)} dBm"))
    if v_trp is not None:
        rows.append(("V-pol TRP", f"{_fmt(v_trp)} dBm"))
    if h_trp is not None and v_trp is not None:
        balance = float(h_trp) - float(v_trp)
        rows.append(("H/V Balance", f"{_fmt(balance, '.1f')} dB"))

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Shading Accent 1"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    _style_header_row(table, brand_dark)

    # Prose assessment
    assessment = f"Total Radiated Power measured at {_fmt(freq, '.0f')} MHz is {_fmt(trp)} dBm"
    if h_trp is not None and v_trp is not None:
        balance = abs(float(h_trp) - float(v_trp))
        if balance < 1:
            assessment += ". H/V polarization power is well-balanced."
        elif balance < 3:
            assessment += f". H/V polarization shows moderate imbalance ({_fmt(balance, '.1f')} dB)."
        else:
            dominant = "H" if float(h_trp) > float(v_trp) else "V"
            assessment += (
                f". Significant H/V polarization imbalance ({_fmt(balance, '.1f')} dB, "
                f"{dominant}-pol dominant).")
    doc.add_paragraph(assessment)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Legacy Table Functions (kept for backward compatibility)
# ---------------------------------------------------------------------------

def _add_gain_stats_table(doc, stats: Dict, brand_dark):
    """Add a formatted gain statistics table to the document."""
    scan_type = stats.get("scan_type", "passive")

    if scan_type == "passive":
        row_data = [
            ("Parameter", "Value"),
            ("Frequency", f"{_fmt(stats.get('frequency_actual'), '.1f')} MHz"),
            ("Peak Gain", f"{_fmt(stats.get('max_gain_dBi'))} dBi"),
            ("Minimum Gain", f"{_fmt(stats.get('min_gain_dBi'))} dBi"),
            ("Average Gain", f"{_fmt(stats.get('avg_gain_dBi'))} dBi"),
            ("Std Deviation", f"{_fmt(stats.get('std_dev_dB'))} dB"),
        ]
        if stats.get("max_hpol_gain_dBi") is not None:
            row_data.append(("H-pol Peak Gain", f"{_fmt(stats.get('max_hpol_gain_dBi'))} dBi"))
        if stats.get("max_vpol_gain_dBi") is not None:
            row_data.append(("V-pol Peak Gain", f"{_fmt(stats.get('max_vpol_gain_dBi'))} dBi"))
    else:
        row_data = [
            ("Parameter", "Value"),
            ("Peak Power", f"{_fmt(stats.get('max_power_dBm'))} dBm"),
            ("Minimum Power", f"{_fmt(stats.get('min_power_dBm'))} dBm"),
            ("Average Power", f"{_fmt(stats.get('avg_power_dBm'))} dBm"),
            ("Std Deviation", f"{_fmt(stats.get('std_dev_dB'))} dB"),
        ]
        if stats.get("TRP_dBm") is not None:
            row_data.append(("TRP", f"{_fmt(stats.get('TRP_dBm'))} dBm"))

    table = doc.add_table(rows=len(row_data), cols=2)
    table.style = "Light Shading Accent 1"

    for i, (label, value) in enumerate(row_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        if i == 0:
            for cell in table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = brand_dark

    doc.add_paragraph()


def _add_freq_comparison_table(doc, antenna_analyzer, add_branded_heading, brand_dark):
    """Add a multi-frequency comparison table (legacy, delegates to new builder)."""
    rep_freqs = _select_representative_frequencies(
        antenna_analyzer.frequencies, 10)
    _build_comparison_table(doc, antenna_analyzer, rep_freqs,
                            brand_dark, add_branded_heading)


# ---------------------------------------------------------------------------
# Main DOCX Builder
# ---------------------------------------------------------------------------

def _build_branded_docx(output_path: str, report_data: Dict,
                        plot_images: Dict[str, List[str]],
                        opts: ReportOptions, provider, metadata: Optional[Dict],
                        measurements: Dict[str, LoadedMeasurement]):
    """Build a professional branded DOCX report with data-driven content.

    Produces engineering-quality sections: data-driven executive summary,
    consolidated performance tables, pattern prose, polarization analysis,
    TRP sections for active data, and data-driven conclusions.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from plot_antenna import config

    # Brand colors
    BRAND_PRIMARY = (
        RGBColor(*config.BRAND_PRIMARY_COLOR)
        if getattr(config, "BRAND_PRIMARY_COLOR", None)
        else RGBColor(70, 130, 180)
    )
    BRAND_DARK = (
        RGBColor(*config.BRAND_DARK_COLOR)
        if getattr(config, "BRAND_DARK_COLOR", None)
        else RGBColor(50, 50, 50)
    )
    BRAND_LIGHT = (
        RGBColor(*config.BRAND_LIGHT_COLOR)
        if getattr(config, "BRAND_LIGHT_COLOR", None)
        else RGBColor(128, 128, 128)
    )

    brand_name = getattr(config, "BRAND_NAME", None)
    brand_tagline = getattr(config, "BRAND_TAGLINE", None)
    brand_website = getattr(config, "BRAND_WEBSITE", None)
    report_subtitle = getattr(config, "REPORT_SUBTITLE", "Antenna Measurement & Analysis Report")

    def add_branded_heading(doc, text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = BRAND_DARK
            if level == 1:
                run.font.size = Pt(getattr(config, "HEADING1_FONT_SIZE", 18))
            else:
                run.font.size = Pt(getattr(config, "HEADING2_FONT_SIZE", 14))
            run.font.bold = True
        return heading

    doc = Document()

    # Pre-compute band info per measurement
    band_info_map: Dict[str, Optional[Dict]] = {}
    for meas_name in report_data.get("measurements", []):
        m = measurements.get(meas_name)
        if m:
            meas_freqs = _filter_frequencies(m.frequencies, opts)
            band_info_map[meas_name] = _detect_rf_band(meas_freqs)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Logo in header
    logo_path = opts.logo_path
    if not logo_path and getattr(config, "LOGO_FILENAME", None):
        candidates = [
            os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                         "assets", config.LOGO_FILENAME),
        ]
        for p in candidates:
            if os.path.exists(p):
                logo_path = p
                break

    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_run = header_para.add_run()
        logo_width = getattr(config, "LOGO_WIDTH_INCHES", 2.0)
        header_run.add_picture(logo_path, width=Inches(logo_width))
        logo_align = getattr(config, "LOGO_ALIGNMENT", "LEFT")
        align_map = {"CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                     "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT}
        header_para.alignment = align_map.get(logo_align, WD_PARAGRAPH_ALIGNMENT.LEFT)

    # ------------------------------------------------------------------ #
    # SECTION 1: Title Page
    # ------------------------------------------------------------------ #
    if opts.include_cover_page:
        title_text = (metadata or {}).get("title", "Antenna Radiation Pattern Test Report")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_heading(title_text, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = BRAND_DARK
        title_run.font.size = Pt(getattr(config, "TITLE_FONT_SIZE", 28))
        title_run.font.bold = True

        if report_subtitle:
            sub = doc.add_paragraph()
            sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            sub_run = sub.add_run(report_subtitle)
            sub_run.font.color.rgb = BRAND_LIGHT
            sub_run.font.size = Pt(getattr(config, "SUBTITLE_FONT_SIZE", 14))
            sub_run.italic = True

        doc.add_paragraph()
        doc.add_paragraph()

        if metadata:
            meta_fields = [
                ("Project:", metadata.get("project_name")),
                ("Antenna Type:", metadata.get("antenna_type")),
                ("Frequency Range:", metadata.get("frequency_range")),
                ("Date:", metadata.get("date", datetime.now().strftime("%B %d, %Y"))),
                ("Prepared by:", metadata.get("author")),
            ]
            for label, value in meta_fields:
                if value:
                    mp = doc.add_paragraph()
                    mp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    lr = mp.add_run(f"{label} ")
                    lr.font.bold = True
                    lr.font.color.rgb = BRAND_DARK
                    lr.font.size = Pt(12)
                    vr = mp.add_run(value)
                    vr.font.color.rgb = BRAND_LIGHT
                    vr.font.size = Pt(12)
                    mp.paragraph_format.space_after = Pt(3)
        else:
            dp = doc.add_paragraph()
            dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            dr = dp.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            dr.font.color.rgb = BRAND_LIGHT
            dr.font.size = Pt(12)

        if brand_website:
            doc.add_paragraph()
            doc.add_paragraph()
            wp = doc.add_paragraph()
            wp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            wr = wp.add_run(brand_website)
            wr.font.color.rgb = BRAND_LIGHT
            wr.font.size = Pt(10)
            wr.italic = True

        doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION 2: Table of Contents
    # ------------------------------------------------------------------ #
    if opts.include_table_of_contents:
        add_branded_heading(doc, "Table of Contents", level=1)
        toc = doc.add_paragraph(
            "< Table of Contents will be generated when document is opened in Word >"
        )
        toc.paragraph_format.space_before = Pt(6)
        doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION 3: Executive Summary (data-driven, AI optional enhancement)
    # ------------------------------------------------------------------ #
    add_branded_heading(doc, "Executive Summary", level=1)

    # Always build data-driven summary as the baseline
    summary_paragraphs = _build_executive_summary(measurements, opts,
                                                   band_info_map=band_info_map)

    if opts.ai_executive_summary:
        # Try AI enhancement; fall back to data-driven if unavailable
        ai_summary = _generate_ai_text(
            provider,
            "You are an RF engineer analyzing antenna test data.\n"
            "Write a concise executive summary (2-3 paragraphs) highlighting "
            "key performance characteristics, any concerns, and overall assessment.",
            report_data, opts, max_tokens=500,
        )
        if ai_summary:
            for para_text in ai_summary.split("\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        else:
            for para in summary_paragraphs:
                doc.add_paragraph(para)
    else:
        for para in summary_paragraphs:
            doc.add_paragraph(para)
    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION 4: Test Configuration (with methodology details)
    # ------------------------------------------------------------------ #
    add_branded_heading(doc, "Test Configuration", level=1)
    doc.add_paragraph(f"Measurements included: {len(report_data['measurements'])}")

    for meas_name in report_data["measurements"]:
        m = measurements.get(meas_name)
        if not m:
            continue

        add_branded_heading(doc, meas_name, level=2)
        test_config = _get_test_configuration(m)

        config_lines = [
            f"Scan type: {test_config['scan_type']}",
            f"Frequencies: {test_config.get('num_frequencies', 'N/A')} points, "
            f"{test_config.get('freq_range', 'N/A')}",
        ]
        if "freq_step" in test_config:
            config_lines.append(f"Frequency step: {test_config['freq_step']}")
        if "theta_range" in test_config:
            config_lines.append(
                f"Theta: {test_config['theta_range']} "
                f"({test_config.get('theta_points', '?')} points"
                f"{', step ' + test_config['theta_step'] if 'theta_step' in test_config else ''})")
        if "phi_range" in test_config:
            config_lines.append(
                f"Phi: {test_config['phi_range']} "
                f"({test_config.get('phi_points', '?')} points"
                f"{', step ' + test_config['phi_step'] if 'phi_step' in test_config else ''})")
        if "spatial_points" in test_config:
            config_lines.append(f"Total spatial points: {test_config['spatial_points']}")

        for line in config_lines:
            doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()

    # ------------------------------------------------------------------ #
    # SECTION 5: Measurement Analysis (consolidated tables)
    # ------------------------------------------------------------------ #
    if opts.include_gain_tables:
        add_branded_heading(doc, "Measurement Analysis", level=1)

        for meas_name in report_data["measurements"]:
            m = measurements.get(meas_name)
            if not m:
                continue

            analyzer, _, err = _get_analyzer_for_measurement(meas_name)
            if analyzer is None:
                continue

            add_branded_heading(doc, f"Performance Summary - {meas_name}", level=2)

            freqs = _filter_frequencies(m.frequencies, opts)
            rep_freqs = _select_representative_frequencies(
                freqs, opts.max_frequencies_in_table,
                band_info=band_info_map.get(meas_name))

            # ONE consolidated table per measurement
            _build_consolidated_performance_table(
                doc, analyzer, meas_name, rep_freqs,
                BRAND_DARK, m.scan_type)

            # Multi-frequency comparison (representative freqs only)
            if m.scan_type == "passive" and len(rep_freqs) >= 2:
                _build_comparison_table(
                    doc, analyzer, rep_freqs, BRAND_DARK, add_branded_heading)

        doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION 6: TRP Analysis (active measurements only)
    # ------------------------------------------------------------------ #
    active_names = [
        n for n in report_data["measurements"]
        if n in measurements and measurements[n].scan_type == "active"
    ]
    if active_names:
        add_branded_heading(doc, "TRP Analysis", level=1)
        for meas_name in active_names:
            m = measurements[meas_name]
            add_branded_heading(doc, meas_name, level=2)
            _build_trp_section(doc, m, BRAND_DARK, add_branded_heading)

    # ------------------------------------------------------------------ #
    # SECTION 7: Measurement Results (embedded plots)
    # ------------------------------------------------------------------ #
    total_images = sum(len(imgs) for imgs in plot_images.values())
    if total_images > 0:
        add_branded_heading(doc, "Measurement Results", level=1)
        figure_num = 1

        for meas_name, img_paths in plot_images.items():
            if not img_paths:
                continue
            add_branded_heading(doc, meas_name, level=2)

            for img_path in _sort_images_by_frequency(img_paths):
                if not os.path.exists(img_path):
                    continue
                doc.add_picture(img_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                caption = doc.add_paragraph()
                cap_run = caption.add_run(f"Figure {figure_num}: ")
                cap_run.bold = True
                cap_run.font.color.rgb = BRAND_DARK
                cap_run.font.size = Pt(11)

                fname_run = caption.add_run(_pretty_caption(os.path.basename(img_path)))
                fname_run.font.color.rgb = BRAND_LIGHT
                fname_run.font.size = Pt(11)
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                figure_num += 1

        doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION 8: Pattern Analysis (prose, not raw dump)
    # ------------------------------------------------------------------ #
    passive_measurements = [
        n for n in report_data["measurements"]
        if n in measurements and measurements[n].scan_type == "passive"
    ]
    if passive_measurements:
        add_branded_heading(doc, "Pattern Analysis", level=1)

        for meas_name in passive_measurements:
            m = measurements[meas_name]
            analyzer, _, err = _get_analyzer_for_measurement(meas_name)
            if analyzer is None:
                continue

            meas_band = band_info_map.get(meas_name)
            meas_freqs = _filter_frequencies(m.frequencies, opts)
            rep_freqs = _select_representative_frequencies(
                meas_freqs, opts.max_plot_frequencies,
                band_info=meas_band)

            add_branded_heading(doc, f"Pattern Analysis - {meas_name}", level=2)

            for freq in rep_freqs:
                prose = _build_pattern_prose(analyzer, freq, meas_name,
                                            band_info=meas_band)
                doc.add_paragraph(prose)

        # Optional AI commentary
        if opts.ai_section_analysis:
            ai_commentary = _generate_ai_text(
                provider,
                "Analyze the radiation patterns and comment on pattern classification, "
                "beamwidth characteristics, front-to-back ratio, and any anomalies.",
                report_data, opts, max_tokens=400,
            )
            if ai_commentary:
                doc.add_paragraph()
                for line in ai_commentary.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

    # ------------------------------------------------------------------ #
    # SECTION 9: Polarization Analysis (consolidated table + summary)
    # ------------------------------------------------------------------ #
    if passive_measurements:
        add_branded_heading(doc, "Polarization Analysis", level=1)

        for meas_name in passive_measurements:
            m = measurements[meas_name]
            analyzer, _, err = _get_analyzer_for_measurement(meas_name)
            if analyzer is None:
                continue

            meas_freqs = _filter_frequencies(m.frequencies, opts)
            rep_freqs = _select_representative_frequencies(
                meas_freqs, opts.max_plot_frequencies,
                band_info=band_info_map.get(meas_name))

            add_branded_heading(doc, f"Polarization - {meas_name}", level=2)
            _build_polarization_table(doc, analyzer, rep_freqs, BRAND_DARK)

            # Summary prose
            if rep_freqs:
                mid_freq = rep_freqs[len(rep_freqs) // 2]
                pol = analyzer.compare_polarizations(frequency=mid_freq)
                note = pol.get("polarization_note")
                balance = pol.get("polarization_balance_dB")
                dominant = pol.get("dominant_pol", "N/A")
                summary_text = (
                    f"At {_fmt(mid_freq, '.0f')} MHz: {dominant}-pol dominant"
                    f" with {_fmt(abs(balance) if balance else None, '.1f')} dB balance")
                if note:
                    summary_text += f". {note}."
                doc.add_paragraph(summary_text)

    # ------------------------------------------------------------------ #
    # SECTION 10: Conclusions and Recommendations (data-driven)
    # ------------------------------------------------------------------ #
    doc.add_page_break()
    add_branded_heading(doc, "Conclusions and Recommendations", level=1)

    if opts.ai_recommendations:
        ai_conclusions = _generate_ai_text(
            provider,
            "Based on the measurements, provide 4-6 bullet-point conclusions and "
            "recommendations for the antenna design. Be specific and actionable.",
            report_data, opts, max_tokens=500,
        )
        if ai_conclusions:
            doc.add_paragraph("Based on the measurement results presented in this report:")
            for line in ai_conclusions.split("\n"):
                line = line.strip().lstrip("-").lstrip("*").lstrip()
                if line:
                    doc.add_paragraph(f"{line}", style="List Bullet")
        else:
            _add_data_driven_conclusions(doc, measurements, opts, band_info_map)
    else:
        _add_data_driven_conclusions(doc, measurements, opts, band_info_map)

    # Brand footer
    if brand_tagline or brand_website:
        doc.add_paragraph()
        if brand_tagline and brand_name:
            fp = doc.add_paragraph()
            fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            fr = fp.add_run(f"{brand_name} | {brand_tagline}")
            fr.font.color.rgb = BRAND_LIGHT
            fr.font.size = Pt(10)
            fr.italic = True
        if brand_website:
            wp = doc.add_paragraph()
            wp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            wr = wp.add_run(brand_website)
            wr.font.color.rgb = BRAND_PRIMARY
            wr.font.size = Pt(10)
            wr.bold = True

    doc.save(output_path)


def _add_data_driven_conclusions(doc, measurements, opts,
                                 band_info_map=None):
    """Add data-driven conclusion bullets."""
    bullets = _build_data_driven_conclusions(measurements, opts,
                                             band_info_map=band_info_map)
    doc.add_paragraph("Based on the measurement results presented in this report:")
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")


def _add_fallback_conclusions(doc):
    """Add generic conclusion bullets when AI is not available."""
    doc.add_paragraph("Based on the measurement results presented in this report:")
    doc.add_paragraph(
        "Review all performance metrics against specification requirements",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Verify antenna performance meets application needs across the operational bandwidth",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Consider additional measurements or design iterations if performance gaps are identified",
        style="List Bullet",
    )


# ---------------------------------------------------------------------------
# Data Preparation
# ---------------------------------------------------------------------------

def _prepare_report_data(measurements: Dict[str, LoadedMeasurement],
                         opts: ReportOptions) -> Dict:
    """Prepare data for report generation based on options."""
    data: Dict[str, Any] = {
        "measurements": [],
        "frequencies": [],
    }

    for name, m in measurements.items():
        if opts.measurements is None or name in opts.measurements:
            data["measurements"].append(name)
            for freq in m.frequencies:
                if opts.frequencies is None or freq in opts.frequencies:
                    if freq not in data["frequencies"]:
                        data["frequencies"].append(freq)

    data["frequencies"].sort()
    return data


# ---------------------------------------------------------------------------
# MCP Tool Registration
# ---------------------------------------------------------------------------

def register_report_tools(mcp):
    """Register report generation tools with the MCP server."""

    @mcp.tool()
    def get_report_options() -> str:
        """
        Get available report configuration options.

        Returns:
            Documentation of all available filtering and configuration options.
        """
        return """
REPORT FILTERING OPTIONS
========================

Use these options with generate_report() to control what's included.

CONTENT FILTERING:
- frequencies: List of frequencies to include (MHz), or null for all
  Example: [2400, 2450, 2500]

- polarizations: Which polarization data to include
  Options: ["total"], ["hpol", "vpol"], ["total", "hpol", "vpol"]
  Default: ["total"]

- measurements: Specific measurement files to include, or null for all

PLOT FILTERING (manages complexity):
- include_2d_plots: true/false (default: true)
  Includes 2D azimuth/elevation pattern cuts

- include_3d_plots: true/false (default: FALSE)
  3D surface plots are large - disabled by default

- max_plot_frequencies: number (default: 5)
  Max per-frequency plots (azimuth cuts, pattern/polarization analysis).
  Selects evenly-spaced representative frequencies (start, end, middle).
  Set higher for more detail, lower for a compact report.

DATA FILTERING:
- include_gain_tables: true/false (default: true)
  Summary gain tables per frequency

- max_frequencies_in_table: number (default: 10)
  Limits table rows for readability

AI CONTENT:
- ai_executive_summary: true/false (default: true)
  AI-generated executive summary

- ai_section_analysis: true/false (default: true)
  AI commentary on each section

- ai_recommendations: true/false (default: true)
  AI-generated design recommendations

- ai_model: "gpt-4o-mini", "gpt-4o", "o3", etc.
  Default: "gpt-4o-mini" (cost-effective)

OUTPUT:
- include_cover_page: true/false (default: true)
- include_table_of_contents: true/false (default: true)

METADATA (optional dict):
- title: Custom report title
- project_name: Project name for cover page
- antenna_type: Antenna type description
- frequency_range: Frequency range string
- author: Report author name
- date: Report date string

EXAMPLE - Minimal Report:
{
    "frequencies": [2450],
    "polarizations": ["total"],
    "include_2d_plots": true,
    "include_3d_plots": false,
    "ai_executive_summary": false,
    "ai_section_analysis": false
}

EXAMPLE - Full Report:
{
    "frequencies": null,
    "polarizations": ["total", "hpol", "vpol"],
    "include_2d_plots": true,
    "include_3d_plots": true,
    "include_gain_tables": true,
    "ai_executive_summary": true,
    "ai_section_analysis": true,
    "ai_recommendations": true,
    "metadata": {
        "title": "BLE Antenna Test Report",
        "project_name": "Product X",
        "antenna_type": "PCB Trace Antenna"
    }
}
"""

    @mcp.tool()
    def generate_report(
        output_path: str,
        options: Optional[Dict[str, Any]] = None,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, str]] = None,
    ) -> str:
        """
        Generate a professional branded antenna test report (DOCX).

        Produces a DOCX with cover page, gain tables, embedded 2D/3D plots,
        AI-generated summaries and conclusions, and branded formatting.

        Args:
            output_path: Path for the output DOCX file
            options: Report options (see get_report_options for details)
            title: Custom report title (default: "Antenna Radiation Pattern Test Report")
            metadata: Dict with project_name, antenna_type, frequency_range, author, date

        Returns:
            Path to generated report and summary of contents.
        """
        measurements = get_loaded_measurements()

        if not measurements:
            return "No data loaded. Use import_antenna_file or import_passive_pair first."

        # Parse options
        opts = ReportOptions()
        if options:
            for key, value in options.items():
                if hasattr(opts, key):
                    setattr(opts, key, value)

        # Merge title into metadata
        if metadata is None:
            metadata = {}
        if title:
            metadata.setdefault("title", title)

        try:
            # 1. Prepare report data
            report_data = _prepare_report_data(measurements, opts)

            if not report_data["measurements"]:
                return "No measurements match the specified filters."

            # 2. Generate plots in temp dir
            plot_images: Dict[str, List[str]] = {}
            temp_dir = tempfile.mkdtemp(prefix="rflect_report_")

            if opts.include_2d_plots or opts.include_3d_plots:
                plot_images = _generate_plots(measurements, opts, temp_dir)

            # 3. Create AI provider (optional)
            provider = _create_llm_provider(opts)

            # 4. Build branded DOCX
            _build_branded_docx(
                output_path, report_data, plot_images,
                opts, provider, metadata, measurements,
            )

            # 5. Summary
            total_plots = sum(len(imgs) for imgs in plot_images.values())
            summary = f"Report generated: {output_path}\n\n"
            summary += "Contents:\n"
            summary += f"- Measurements: {len(report_data['measurements'])}\n"
            summary += f"- Frequencies: {report_data['frequencies']}\n"
            summary += f"- Embedded plots: {total_plots}\n"
            summary += f"- Gain tables: {'Yes' if opts.include_gain_tables else 'No'}\n"
            summary += f"- AI provider: {'Connected' if provider else 'None (fallback text used)'}\n"
            summary += f"- Cover page: {'Yes' if opts.include_cover_page else 'No'}\n"

            return summary

        except Exception as e:
            return f"Error generating report: {str(e)}"

    @mcp.tool()
    def preview_report(options: Optional[Dict[str, Any]] = None) -> str:
        """
        Preview what would be included in a report without generating it.

        Args:
            options: Same options as generate_report

        Returns:
            Summary of what the report would contain.
        """
        measurements = get_loaded_measurements()

        if not measurements:
            return "No data loaded. Use import_antenna_file or import_passive_pair first."

        # Parse options
        opts = ReportOptions()
        if options:
            for key, value in options.items():
                if hasattr(opts, key):
                    setattr(opts, key, value)

        report_data = _prepare_report_data(measurements, opts)

        preview = "REPORT PREVIEW\n"
        preview += "=" * 40 + "\n\n"

        # Measurements
        preview += f"MEASUREMENTS ({len(report_data['measurements'])})\n"
        for name in report_data["measurements"]:
            m = measurements.get(name)
            scan = m.scan_type if m else "?"
            preview += f"  - {name} ({scan})\n"

        # Frequencies
        preview += f"\nFREQUENCIES ({len(report_data['frequencies'])})\n"
        for freq in report_data["frequencies"][:10]:
            preview += f"  - {freq} MHz\n"
        if len(report_data["frequencies"]) > 10:
            preview += f"  ... and {len(report_data['frequencies']) - 10} more\n"

        # Sections
        preview += "\nSECTIONS\n"
        preview += f"  [{'x' if opts.include_cover_page else ' '}] Cover Page (branded)\n"
        preview += f"  [{'x' if opts.include_table_of_contents else ' '}] Table of Contents\n"
        preview += f"  [{'x' if opts.ai_executive_summary else ' '}] Executive Summary (AI)\n"
        preview += "  [x] Test Configuration\n"
        preview += f"  [{'x' if opts.include_gain_tables else ' '}] Gain Summary Tables\n"
        preview += f"  [{'x' if opts.include_2d_plots else ' '}] 2D Pattern Plots\n"
        preview += f"  [{'x' if opts.include_3d_plots else ' '}] 3D Pattern Plots\n"
        preview += f"  [{'x' if opts.ai_section_analysis else ' '}] Pattern Analysis (AI)\n"
        preview += f"  [{'x' if opts.ai_recommendations else ' '}] Conclusions (AI)\n"

        # Estimate plot count (respects max_plot_frequencies)
        n_meas = len(report_data["measurements"])
        n_plot_freqs = min(len(report_data["frequencies"]), opts.max_plot_frequencies)
        plot_count = 0
        if opts.include_2d_plots:
            # Per-freq azimuth cuts + 3 summary plots per measurement
            plot_count += n_plot_freqs * n_meas + 3 * n_meas
        if opts.include_3d_plots:
            # One set of 3D plots per measurement (at center freq)
            plot_count += n_meas * len(opts.polarizations) * 2

        preview += f"\nESTIMATED COMPLEXITY\n"
        preview += f"  Plots: ~{plot_count}\n"
        n_freqs = len(report_data["frequencies"])
        preview += f"  Tables: ~{n_freqs if opts.include_gain_tables else 0}\n"
        ai_count = sum([opts.ai_executive_summary, opts.ai_section_analysis, opts.ai_recommendations])
        preview += f"  AI Sections: {ai_count}\n"

        if plot_count > 20:
            preview += f"\nWarning: {plot_count} plots may make the report very large.\n"
            preview += "Consider reducing frequencies or disabling 3D plots.\n"

        return preview
